import os
import io
from google.cloud import vision
from google.oauth2 import service_account
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import argparse

class OCRToWord:
    def __init__(self, credentials_path=None):
        """初始化OCR客户端"""
        if credentials_path and os.path.exists(credentials_path):
            credentials = service_account.Credentials.from_service_account_file(
                credentials_path
            )
            self.client = vision.ImageAnnotatorClient(credentials=credentials)
        else:
            # 尝试使用环境变量中的凭证
            self.client = vision.ImageAnnotatorClient()

        self.document = Document()

    def detect_text_from_image(self, image_path):
        """从图片中提取文本（带段落识别）"""
        with io.open(image_path, 'rb') as image_file:
            content = image_file.read()

        image = vision.Image(content=content)

        # 启用文档文本检测（保持段落结构）
        response = self.client.document_text_detection(image=image)
        texts = response.full_text_annotation

        if response.error.message:
            raise Exception(f'Google Vision API错误: {response.error.message}')

        return texts

    def extract_paragraphs(self, full_text_annotation):
        """从检测结果中提取段落"""
        paragraphs = []
        current_paragraph = []

        for page in full_text_annotation.pages:
            for block in page.blocks:
                # 按块处理文本（通常每个块对应一个段落）
                block_text = []
                for paragraph in block.paragraphs:
                    words = []
                    for word in paragraph.words:
                        word_text = ''.join([symbol.text for symbol in word.symbols])
                        words.append(word_text)
                    block_text.append(' '.join(words))

                paragraph_text = '\n'.join(block_text)
                if paragraph_text.strip():
                    paragraphs.append(paragraph_text)

        return paragraphs

    def save_to_word(self, paragraphs, output_path):
        """将段落保存到Word文档"""
        # 设置文档样式
        style = self.document.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # 添加标题
        self.document.add_heading('OCR识别结果', 0)

        # 添加段落
        for i, paragraph_text in enumerate(paragraphs, 1):
            # 添加段落标题
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 添加文本
            para.add_run(f'段落 {i}:').bold = True
            para.add_run(f'\n{paragraph_text}\n')

        # 保存文档
        self.document.save(output_path)
        print(f"文档已保存到: {output_path}")

    def process_image(self, image_path, output_docx=None):
        """处理单个图片"""
        if output_docx is None:
            output_docx = os.path.splitext(image_path)[0] + '_ocr.docxx'

        print(f"正在处理: {image_path}")

        # 执行OCR
        text_annotation = self.detect_text_from_image(image_path)

        # 提取段落
        paragraphs = self.extract_paragraphs(text_annotation)

        # 保存到Word
        self.save_to_word(paragraphs, output_docx)

        # 打印统计信息
        print(f"识别到 {len(paragraphs)} 个段落")
        print(f"总字符数: {len(text_annotation.text)}")

        return paragraphs

def main():
    parser = argparse.ArgumentParser(description='OCR图片转Word文档')
    parser.add_argument('image_path', help='输入图片路径')
    parser.add_argument('-o', '--output', help='输出Word文档路径')
    parser.add_argument('-c', '--credentials', help='Google API凭证文件路径')

    args = parser.parse_args()

    # 创建OCR处理器
    ocr_processor = OCRToWord(args.credentials)

    # 处理图片
    try:
        paragraphs = ocr_processor.process_image(args.image_path, args.output)

        # 在控制台显示识别结果
        print("\n识别结果预览:")
        print("-" * 50)
        for i, para in enumerate(paragraphs[:3], 1):  # 只显示前3个段落
            print(f"段落 {i}: {para[:100]}..." if len(para) > 100 else f"段落 {i}: {para}")
        if len(paragraphs) > 3:
            print(f"... 还有 {len(paragraphs) - 3} 个段落")

    except Exception as e:
        print(f"处理失败: {e}")

if __name__ == "__main__":
    main()