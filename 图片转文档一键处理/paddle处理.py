# -*- coding: utf-8 -*-
import os

import base64

from PIL import Image
from paddleocr import PaddleOCR

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from openai import OpenAI

from io import BytesIO

# ================= OCR 初始化 =================
ocr = PaddleOCR(
    use_doc_orientation_classify=True, # 通过 use_doc_orientation_classify 参数指定不使用文档方向分类模型
    use_doc_unwarping=True, # 通过 use_doc_unwarping 参数指定不使用文本图像矫正模型
    use_textline_orientation=False, # 通过 use_textline_orientation 参数指定不使用文本行方向分类模型
)
# ocr = PaddleOCR(lang="en") # 通过 lang 参数来使用英文模型
# ocr = PaddleOCR(ocr_version="PP-OCRv4") # 通过 ocr_version 参数来使用 PP-OCR 其他版本
# ocr = PaddleOCR(device="gpu") # 通过 device 参数使得在模型推理时使用 GPU
# ocr = PaddleOCR(
#     text_detection_model_name="PP-OCRv5_server_det",
#     text_recognition_model_name="PP-OCRv5_server_rec",
#     use_doc_orientation_classify=False,
#     use_doc_unwarping=False,
#     use_textline_orientation=False,
# ) # 更换 PP-OCRv5_server 模型
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff"}

import cv2

def resize_image_for_ocr(path, max_side=4000):
    img = cv2.imread(path)
    h, w = img.shape[:2]
    scale = min(max_side / max(h, w), 1.0)
    if scale < 1.0:
        img = cv2.resize(img, (int(w*scale), int(h*scale)))
    return img


def ocr_and_extract_text(image_path):
    paragraphs = []
    buffer = ""

    img = resize_image_for_ocr(image_path, max_side=4000)


    # 图片识别
    result = ocr.predict(img)
    # print(result)
    # 遍历 PaddleOCR 输出
    # result 可能是：[[[bbox, (text, score)], ...], ...]
    for res in result:
        lines = []
        data = res.json
        texts = data.get("res", {}).get("rec_texts", [])
        lines.extend(texts)
        # print(lines)


            # 将行文本合并成段落
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 极短行直接拼接
            if len(line) <= 4:
                buffer += line
                continue

            buffer += line

            # 以句号/问号/感叹号结尾且长度足够 → 分段
            if buffer[-1] in "。！？…" and len(buffer) >= 30:
                paragraphs.append(buffer)
                buffer = ""

        if buffer:
            paragraphs.append(buffer)

    return paragraphs







def process_folder(folder_path, root_dir, out_root, log_callback=print):
    folder_name = os.path.basename(folder_path)

    # ⭐ 关键：复制父文件夹结构
    rel_path = os.path.relpath(folder_path, root_dir)
    out_dir = os.path.join(out_root, rel_path)

    os.makedirs(out_dir, exist_ok=True)
    doc_path = os.path.join(out_dir, f"{folder_name}.docx")

    all_text_blocks = []

    for filename in sorted(os.listdir(folder_path)):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
            img_path = os.path.join(folder_path, filename)
            log_callback(f"正在识别：{img_path}")
            paragraphs = ocr_and_extract_text(img_path)
            all_text_blocks.append(paragraphs)

    if all_text_blocks:








def create_word(doc_path, all_text_blocks,folder_display_name):

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # 修改前
    para_before = doc.add_paragraph("修改前：")
    para_before.paragraph_format.first_line_indent = Cm(0.74)
    para_before.paragraph_format.space_before = Pt(0)
    para_before.paragraph_format.space_after = Pt(0)
    para_before.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    para_before.paragraph_format.line_spacing = Pt(12)

    # 姓名
    para_name = doc.add_paragraph(f"——{folder_display_name}")
    para_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_name.paragraph_format.space_before = Pt(0)
    para_name.paragraph_format.space_after = Pt(0)
    para_name.paragraph_format.line_spacing = Pt(12)

    for para_text in all_text_blocks:
        # OCR 段落
        p = doc.add_paragraph(para_text)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        p.paragraph_format.line_spacing = Pt(12)

    doc.add_page_break()
    # 修改后
    para_after = doc.add_paragraph("修改后：")
    para_after.paragraph_format.first_line_indent = Cm(0.74)
    para_after.paragraph_format.space_before = Pt(0)
    para_after.paragraph_format.space_after = Pt(0)
    para_after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    para_after.paragraph_format.line_spacing = Pt(12)

    doc.save(doc_path)


def has_images(folder_path):
    return any(
        f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
        for f in os.listdir(folder_path)
    )

def process_all(root_dir, out_root, log_callback=print):
    os.makedirs(out_root, exist_ok=True)

    for root, dirs, files in os.walk(root_dir):
        if has_images(root):
            process_folder(root, root_dir, out_root, log_callback)






if __name__ == "__main__":
    img_root = r"/vol1/1000/2T_transfer/project/python/图片转文档一键处理/img/初一下午线下班"
    out_root = r"./docxx"
    process_all(img_root, out_root)
