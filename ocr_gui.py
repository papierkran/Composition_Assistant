import base64
import json
import threading
from concurrent.futures import FIRST_COMPLETED, ThreadPoolExecutor, wait
from datetime import datetime
import sys
import os
import re
from pathlib import Path
from copy import deepcopy

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QCheckBox, QComboBox, QTextEdit,
    QScrollArea, QFrame, QTableWidget, QTableWidgetItem, QHeaderView,
    QAbstractItemView, QSplitter, QGroupBox, QFormLayout, QSizePolicy,
    QFileDialog, QMessageBox,
)
from PySide6.QtCore import Qt, Signal, QObject, QTimer
from PySide6.QtGui import QFont, QColor, QIcon

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from openai import OpenAI


# ================= 默认配置 =================
# 打包后 _MEIPASS 是临时目录，EXE 所在目录才是用户目录
def _exe_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent

PERSONAL_CONFIG = Path(r"D:\person_data\ocer助手\presson.json")
LOCAL_CONFIG = _exe_dir() / "config.json"

DEFAULT_CONFIG = {
    "OCR": {
        "PROVIDER": "xfyun_handwriting",
        "XFYUN": {
            "URL": "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting",
            "APPID": "",
            "API_KEY": "",
            "LANGUAGE": "cn|en",
            "LOCATION": "false",
        },
    },
    "LLM": {
        "PROVIDERS": {
            "deepseek": {"API_KEY": "", "MODEL": "deepseek-chat", "BASE_URL": "https://api.deepseek.com/v1"},
            "openai": {"API_KEY": "", "MODEL": "gpt-4o-mini", "BASE_URL": "https://api.openai.com/v1"},
            "custom": {"API_KEY": "", "MODEL": "", "BASE_URL": ""},
        },
        "TASKS": {
            "typo_fix": {"ENABLED": False, "PROVIDER": "deepseek", "PROMPT": "{text}"},
            "editor": {"ENABLED": False, "PROVIDER": "deepseek", "PROMPT": "{text}", "COUNT_MIN": None, "COUNT_MAX": None},
        },
    },
    "APP": {"ROOT_DIR": "", "DEBUG": False},
}

from config_migrate import ensure_new_schema


def load_config(path: Path = None):
    """加载配置：优先个人目录 → 当前目录 → 默认"""
    if path:
        cfg_path = Path(path)
    elif PERSONAL_CONFIG.exists():
        cfg_path = PERSONAL_CONFIG
    elif LOCAL_CONFIG.exists():
        cfg_path = LOCAL_CONFIG
    else:
        return deepcopy(DEFAULT_CONFIG)
    try:
        with cfg_path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return deepcopy(DEFAULT_CONFIG)


def save_config(config, path: Path = None):
    """保存配置：优先保存到个人目录，不存在则保存到当前目录"""
    if path:
        cfg_path = Path(path)
    elif PERSONAL_CONFIG.parent.exists():
        cfg_path = PERSONAL_CONFIG
    else:
        cfg_path = LOCAL_CONFIG
    cfg_path.parent.mkdir(parents=True, exist_ok=True)
    with cfg_path.open("w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


# ================= Provider helpers =================
API_PROVIDER_PRESETS = {"deepseek": "https://api.deepseek.com/v1", "openai": "https://api.openai.com/v1"}


def _normalize_provider_name(name: str) -> str:
    return (name or "").strip().lower()


def _ensure_provider_exists(config, name: str):
    p_name = _normalize_provider_name(name)
    if not p_name:
        return ""
    config.setdefault("LLM", {}).setdefault("PROVIDERS", {}).setdefault(p_name, {})
    p_cfg = config["LLM"]["PROVIDERS"][p_name]
    p_cfg.setdefault("API_KEY", "")
    p_cfg.setdefault("BASE_URL", API_PROVIDER_PRESETS.get(p_name, ""))
    p_cfg.setdefault("MODEL", "deepseek-chat" if p_name == "deepseek" else "gpt-4o-mini")
    return p_name


def _provider_name_list(config):
    providers = (config.get("LLM", {}) or {}).get("PROVIDERS", {}) or {}
    names = {k for k in providers.keys() if isinstance(k, str) and k.strip()}
    names.update(API_PROVIDER_PRESETS.keys())
    return sorted(names)


# ================= File helpers =================
def iter_files_limited(folder, max_depth=4):
    folder = os.path.abspath(folder)
    for root_dir, dirs, files in os.walk(folder, topdown=True):
        rel = os.path.relpath(root_dir, folder)
        depth = 0 if rel == os.curdir else len(rel.split(os.sep))
        if depth >= max_depth - 1:
            dirs[:] = []
        yield root_dir, files


def has_images_folder(path: str) -> bool:
    try:
        return any(
            f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
            for f in os.listdir(path)
            if os.path.isfile(os.path.join(path, f))
        )
    except Exception:
        return False


def scan_folder_for_tasks(folder: str) -> list[str]:
    folder = os.path.abspath(folder)
    tasks = []
    if has_images_folder(folder):
        tasks.append(folder)
    try:
        for name in sorted(os.listdir(folder)):
            if name == "旧":
                continue
            child = os.path.join(folder, name)
            if os.path.isdir(child) and has_images_folder(child):
                tasks.append(child)
    except Exception:
        pass
    return tasks


def infer_student_and_essay(folder_name: str):
    if "_" in folder_name:
        parts = folder_name.split("_", 1)
        return parts[0], parts[1]
    if "-" in folder_name:
        parts = folder_name.split("-", 1)
        return parts[0], parts[1]
    return folder_name, folder_name


def count_existing_docx_chars(folder_path: str) -> str:
    try:
        doc_name = os.path.basename(folder_path)
        docx_path = os.path.join(folder_path, f"{doc_name}.docx")
        if os.path.isfile(docx_path):
            doc = Document(docx_path)
            total = sum(len(p.text.strip()) for p in doc.paragraphs if p.text.strip())
            return str(total)
    except Exception:
        pass
    return ""


def count_chinese_characters(text: str) -> int:
    return sum(1 for ch in text if not ch.isspace())


def determine_word_count_bounds(original_count: int):
    if original_count >= 850:
        return max(700, original_count - 30), original_count + 30
    if original_count >= 800:
        return 820, 850
    return 700, 820


# ================= Log signal =================
class LogSignal(QObject):
    log_message = Signal(str)
    task_status = Signal(str, str, str, str, str)


# ================= Collapsible Section =================
class CollapsibleSection(QWidget):
    def __init__(self, title="", collapsed=True, parent=None):
        super().__init__(parent)
        self._collapsed = collapsed
        self._title = title

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.toggle_btn = QPushButton(("▶ " + title) if collapsed else ("▼ " + title))
        self.toggle_btn.setStyleSheet(
            "QPushButton { text-align: left; border: none; background: transparent;"
            "color: #1a73e8; font-weight: bold; font-size: 13px; padding: 6px 8px; }"
            "QPushButton:hover { background: #e8f0fe; }"
        )
        self.toggle_btn.clicked.connect(self._toggle)
        layout.addWidget(self.toggle_btn)

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(8, 4, 8, 8)
        layout.addWidget(self.content_widget)

        if collapsed:
            self.content_widget.setVisible(False)

    def _toggle(self):
        self._collapsed = not self._collapsed
        self.content_widget.setVisible(not self._collapsed)
        prefix = "▼ " if not self._collapsed else "▶ "
        self.toggle_btn.setText(prefix + self._title)

    def add_widget(self, w):
        self.content_layout.addWidget(w)

    def add_layout(self, l):
        self.content_layout.addLayout(l)


# ================= Main Window =================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.config = ensure_new_schema(load_config())
        self.hidden_api_keys = {}
        self.task_queue = []
        self.completed_tasks = set()  # 已完成的任务路径集合
        self.finished_tasks = set()  # 当前运行批次中已经尝试过的任务路径集合
        self.in_progress_tasks = set()  # 当前正在处理的任务路径集合
        self.is_processing = False
        self.max_parallel_tasks = 3
        self.queue_lock = threading.Lock()
        self.log_signal = LogSignal()
        self.log_signal.log_message.connect(self._append_log)
        self.log_signal.task_status.connect(self._update_task_status)

        self.setWindowTitle("Composition OCR Assistant 作文修改助手 v1.1")
        self.resize(1100, 800)

        # Set icon
        for ico_name in ("app.ico",):
            ico_path = Path(__file__).resolve().parent / ico_name
            if ico_path.exists():
                self.setWindowIcon(QIcon(str(ico_path)))
                break

        # Central widget
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(4)

        # Top bar
        top_bar = QHBoxLayout()
        top_bar.addWidget(QLabel("功能选择:"))
        self.btn_ocr = QPushButton("图片转作文")
        self.btn_ocr.setFixedWidth(120)
        self.btn_ai = QPushButton("docx作文处理")
        self.btn_ai.setFixedWidth(120)
        self.btn_config = QPushButton("配置编辑")
        self.btn_config.setFixedWidth(100)
        self.btn_config.clicked.connect(self._open_config_editor)
        top_bar.addWidget(self.btn_ocr)
        top_bar.addWidget(self.btn_ai)
        top_bar.addStretch()
        top_bar.addWidget(self.btn_config)
        main_layout.addLayout(top_bar)

        # Page stack
        self.page_ocr = QWidget()
        self.page_ai = QWidget()
        self.page_ocr.hide()
        self.page_ai.hide()
        main_layout.addWidget(self.page_ocr)
        main_layout.addWidget(self.page_ai)

        self.btn_ocr.clicked.connect(lambda: self._show_page("ocr"))
        self.btn_ai.clicked.connect(lambda: self._show_page("ai"))

        self._init_page_ocr()
        self._init_page_ai()
        self._show_page("ocr")

    def _show_page(self, name):
        self.page_ocr.hide()
        self.page_ai.hide()
        if name == "ocr":
            self.page_ocr.show()
        else:
            self.page_ai.show()

    def _append_log(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.log_text.append(f"[{ts}] {msg}")

    def _append_log_ai(self, msg):
        ts = datetime.now().strftime("%H:%M:%S")
        self.ai_log_text.append(f"[{ts}] {msg}")

    def _open_config_editor(self):
        from config_editor_ui import open_config_editor_form
        # 确定实际配置文件路径
        if PERSONAL_CONFIG.exists():
            cfg_file = PERSONAL_CONFIG
        elif LOCAL_CONFIG.exists():
            cfg_file = LOCAL_CONFIG
        else:
            cfg_file = PERSONAL_CONFIG  # 默认保存到个人目录
        open_config_editor_form(
            parent=self,
            config=self.config,
            config_file=cfg_file,
            hidden_api_keys=self.hidden_api_keys,
            on_saved=self._on_config_saved,
        )

    def _on_config_saved(self, new_cfg):
        self.config = ensure_new_schema(new_cfg)
        # Refresh OCR fields
        self.url_entry.setText(self.config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
        self.appid_entry.setText(self.config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
        self.path_entry.setText(self.config.get("APP", {}).get("ROOT_DIR", ""))
        # Refresh checkboxes
        self.use_typo_fix.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
        self.typo_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT", "{text}"))
        self.use_editor.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
        self.editor_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"))
        # Refresh provider combos
        names = _provider_name_list(self.config)
        for combo in [self.typo_provider_combo, self.editor_provider_combo]:
            current = combo.currentText()
            combo.clear()
            combo.addItems(names)
            idx = combo.findText(current)
            if idx >= 0:
                combo.setCurrentIndex(idx)

    # ===================== PAGE 1: OCR =====================
    def _init_page_ocr(self):
        layout = QVBoxLayout(self.page_ocr)
        layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(4, 4, 4, 4)
        scroll_layout.setSpacing(6)

        # ---- Baidu correction (collapsible) ----
        baidu_sec = CollapsibleSection("百度图片矫正（OCR前自动矫正倾斜/弯曲文档）", collapsed=True)
        self.use_baidu_correction = QCheckBox("启用图片矫正（去阴影+透视变换）")
        self.use_baidu_correction.setChecked(bool(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED", False)))
        baidu_sec.add_widget(self.use_baidu_correction)

        baidu_key_layout = QHBoxLayout()
        baidu_key_layout.addWidget(QLabel("百度 API Key"))
        self.baidu_api_key_entry = QLineEdit(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("API_KEY", ""))
        baidu_key_layout.addWidget(self.baidu_api_key_entry, 1)
        baidu_sec.add_layout(baidu_key_layout)

        baidu_secret_layout = QHBoxLayout()
        baidu_secret_layout.addWidget(QLabel("百度 Secret Key"))
        self.baidu_secret_key_entry = QLineEdit(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY", ""))
        baidu_secret_key_entry = QLineEdit(self.config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY", ""))
        baidu_secret_layout.addWidget(baidu_secret_key_entry, 1)
        baidu_sec.add_layout(baidu_secret_layout)
        self.baidu_secret_key_entry = baidu_secret_key_entry

        scroll_layout.addWidget(baidu_sec)

        # ---- OCR config (collapsible) ----
        ocr_sec = CollapsibleSection("OCR 识别配置", collapsed=True)

        ocr_url_layout = QHBoxLayout()
        ocr_url_layout.addWidget(QLabel("OCR 接口 URL"))
        self.url_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
        ocr_url_layout.addWidget(self.url_entry, 1)
        ocr_sec.add_layout(ocr_url_layout)

        ocr_appid_layout = QHBoxLayout()
        ocr_appid_layout.addWidget(QLabel("APPID"))
        self.appid_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
        ocr_appid_layout.addWidget(self.appid_entry, 1)
        ocr_sec.add_layout(ocr_appid_layout)

        ocr_apikey_layout = QHBoxLayout()
        ocr_apikey_layout.addWidget(QLabel("API_KEY"))
        self.apikey_entry = QLineEdit(self.config.get("OCR", {}).get("XFYUN", {}).get("API_KEY", ""))
        ocr_apikey_layout.addWidget(self.apikey_entry, 1)
        ocr_sec.add_layout(ocr_apikey_layout)

        scroll_layout.addWidget(ocr_sec)

        # ---- AI typo fix ----
        typo_group = QGroupBox("第一步：AI 错别字修正")
        typo_layout = QVBoxLayout()

        # API Key
        row1 = QHBoxLayout()
        row1.addWidget(QLabel("AI API Key"))
        self.typo_api_key = QLineEdit()
        self.typo_api_key.setEchoMode(QLineEdit.Password)
        typo_provider = _normalize_provider_name(
            (self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROVIDER", "deepseek")
        ) or "deepseek"
        self.typo_api_key.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(typo_provider, {}) or {}).get("API_KEY", ""))
        row1.addWidget(self.typo_api_key, 1)
        self.use_typo_fix = QCheckBox("启用 AI 错别字自动修正（较慢）")
        self.use_typo_fix.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
        row1.addWidget(self.use_typo_fix)
        typo_layout.addLayout(row1)

        # Provider + Base URL
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("AI Provider"))
        self.typo_provider_combo = QComboBox()
        provider_names = _provider_name_list(self.config)
        self.typo_provider_combo.addItems(provider_names)
        idx = self.typo_provider_combo.findText(typo_provider)
        if idx >= 0:
            self.typo_provider_combo.setCurrentIndex(idx)
        self.typo_provider_combo.currentTextChanged.connect(self._on_typo_provider_change)
        row2.addWidget(self.typo_provider_combo)
        btn_new_typo = QPushButton("+ 新增")
        btn_new_typo.setFixedWidth(68)
        btn_new_typo.clicked.connect(lambda: self._add_provider(self.typo_provider_combo))
        row2.addWidget(btn_new_typo)
        row2.addWidget(QLabel("Base URL"))
        self.typo_base_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(typo_provider, {}) or {}).get("BASE_URL", ""))
        row2.addWidget(self.typo_base_entry, 1)
        typo_layout.addLayout(row2)

        # Prompt
        row3 = QHBoxLayout()
        row3.addWidget(QLabel("自定义提示词"))
        self.typo_prompt_text = QTextEdit()
        self.typo_prompt_text.setMaximumHeight(120)
        self.typo_prompt_text.setPlainText(
            (self.config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT")
            or DEFAULT_CONFIG["LLM"]["TASKS"]["typo_fix"]["PROMPT"]
        )
        row3.addWidget(self.typo_prompt_text, 1)
        typo_layout.addLayout(row3)

        typo_group.setLayout(typo_layout)
        scroll_layout.addWidget(typo_group)

        # ---- Editor (step 2) ----
        editor_group = QGroupBox("第二步：AI 修改作文")
        editor_layout = QVBoxLayout()

        row4 = QHBoxLayout()
        self.use_editor = QCheckBox("启用 第二步 修改作文")
        self.use_editor.setChecked(bool((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
        row4.addWidget(self.use_editor)
        editor_layout.addLayout(row4)

        row5 = QHBoxLayout()
        row5.addWidget(QLabel("AI API Key"))
        self.editor_api_key = QLineEdit()
        self.editor_api_key.setEchoMode(QLineEdit.Password)
        editor_provider = _normalize_provider_name(
            (self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")
        ) or "deepseek"
        self.editor_api_key.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(editor_provider, {}) or {}).get("API_KEY", ""))
        row5.addWidget(self.editor_api_key, 1)
        editor_layout.addLayout(row5)

        row6 = QHBoxLayout()
        row6.addWidget(QLabel("AI Provider"))
        self.editor_provider_combo = QComboBox()
        self.editor_provider_combo.addItems(provider_names)
        idx = self.editor_provider_combo.findText(editor_provider)
        if idx >= 0:
            self.editor_provider_combo.setCurrentIndex(idx)
        self.editor_provider_combo.currentTextChanged.connect(self._on_editor_provider_change)
        row6.addWidget(self.editor_provider_combo)
        btn_new_editor = QPushButton("+ 新增")
        btn_new_editor.setFixedWidth(68)
        btn_new_editor.clicked.connect(lambda: self._add_provider(self.editor_provider_combo))
        row6.addWidget(btn_new_editor)
        row6.addWidget(QLabel("Base URL"))
        self.editor_base_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(editor_provider, {}) or {}).get("BASE_URL", ""))
        row6.addWidget(self.editor_base_entry, 1)
        editor_layout.addLayout(row6)

        row7 = QHBoxLayout()
        row7.addWidget(QLabel("自定义提示词"))
        self.editor_prompt_text = QTextEdit()
        self.editor_prompt_text.setMaximumHeight(120)
        self.editor_prompt_text.setPlainText(
            (self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT")
            or DEFAULT_CONFIG["LLM"]["TASKS"]["editor"]["PROMPT"]
        )
        row7.addWidget(self.editor_prompt_text, 1)
        editor_layout.addLayout(row7)

        row8 = QHBoxLayout()
        row8.addWidget(QLabel("目标字数"))
        self.editor_count_min = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""))
        self.editor_count_min.setFixedWidth(100)
        row8.addWidget(self.editor_count_min)
        row8.addWidget(QLabel("-"))
        self.editor_count_max = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""))
        self.editor_count_max.setFixedWidth(100)
        row8.addWidget(self.editor_count_max)
        row8.addWidget(QLabel("（空白表示自动）"))
        row8.addStretch()
        editor_layout.addLayout(row8)

        editor_group.setLayout(editor_layout)
        scroll_layout.addWidget(editor_group)

        # ---- Path & Start ----
        path_group = QGroupBox("文件路径与任务")
        path_layout = QVBoxLayout()

        path_row = QHBoxLayout()
        path_row.addWidget(QLabel("作文文件夹路径"))
        self.path_entry = QLineEdit(self.config["APP"]["ROOT_DIR"])
        path_row.addWidget(self.path_entry, 1)
        btn_browse = QPushButton("浏览")
        btn_browse.setFixedWidth(70)
        btn_browse.clicked.connect(self._browse_folder)
        path_row.addWidget(btn_browse)
        path_layout.addLayout(path_row)

        # Start button (放在路径下方、任务列表前)
        btn_start = QPushButton("🚀 开始处理（自动读取路径下任务并开始）")
        btn_start.setStyleSheet("background-color: #4CAF50; color: white; font-size: 14px; padding: 8px;")
        btn_start.clicked.connect(self._start_processing)
        path_layout.addWidget(btn_start)

        path_group.setLayout(path_layout)
        scroll_layout.addWidget(path_group)

        # Task queue table (任务日志合一)
        self.queue_table = QTableWidget()
        self.queue_table.setColumnCount(9)
        self.queue_table.setHorizontalHeaderLabels(["序号", "学生姓名", "文件路径", "作文名称", "修改前字数", "当前步骤", "修改后字数", "状态", "实时日志"])
        header_view = self.queue_table.horizontalHeader()
        header_view.setSectionResizeMode(0, QHeaderView.Fixed)
        header_view.setSectionResizeMode(1, QHeaderView.Fixed)
        header_view.setSectionResizeMode(2, QHeaderView.Stretch)
        header_view.setSectionResizeMode(3, QHeaderView.Fixed)
        header_view.setSectionResizeMode(4, QHeaderView.Fixed)
        header_view.setSectionResizeMode(5, QHeaderView.Fixed)
        header_view.setSectionResizeMode(6, QHeaderView.Fixed)
        header_view.setSectionResizeMode(7, QHeaderView.Fixed)
        header_view.setSectionResizeMode(8, QHeaderView.Stretch)
        self.queue_table.setColumnWidth(0, 40)
        self.queue_table.setColumnWidth(1, 100)
        self.queue_table.setColumnWidth(3, 120)
        self.queue_table.setColumnWidth(4, 80)
        self.queue_table.setColumnWidth(5, 100)
        self.queue_table.setColumnWidth(6, 80)
        self.queue_table.setColumnWidth(7, 70)
        self.queue_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.queue_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.queue_table.verticalHeader().setVisible(False)
        self.queue_table.setMinimumHeight(280)
        self.queue_table.setMaximumHeight(400)
        scroll_layout.addWidget(self.queue_table)

        queue_btns = QHBoxLayout()
        btn_add_task = QPushButton("添加")
        btn_add_task.clicked.connect(self._add_task)
        btn_del_task = QPushButton("删除")
        btn_del_task.clicked.connect(self._remove_task)
        btn_requeue_task = QPushButton("重新加入")
        btn_requeue_task.clicked.connect(self._requeue_selected_tasks)
        btn_load_task = QPushButton("读取")
        btn_load_task.clicked.connect(self._load_tasks)
        btn_refresh_task = QPushButton("刷新队列")
        btn_refresh_task.clicked.connect(self._refresh_queue)
        queue_btns.addWidget(btn_add_task)
        queue_btns.addWidget(btn_del_task)
        queue_btns.addWidget(btn_requeue_task)
        queue_btns.addStretch()
        queue_btns.addWidget(btn_load_task)
        queue_btns.addWidget(btn_refresh_task)
        scroll_layout.addLayout(queue_btns)

        # Log (collapsible)
        self.log_section = CollapsibleSection("运行日志", collapsed=True)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(360)
        self.log_section.add_widget(self.log_text)
        scroll_layout.addWidget(self.log_section)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

    def _on_typo_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        self.typo_base_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

    def _on_editor_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        self.editor_base_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

    def _add_provider(self, combo):
        from PySide6.QtWidgets import QInputDialog
        name, ok = QInputDialog.getText(self, "新增 AI Provider", "Provider 名称（如 xai / moonshot）:")
        if ok and name.strip():
            p_name = _normalize_provider_exists(self.config, name.strip())
            _ensure_provider_exists(self.config, p_name)
            save_config(self.config)
            names = _provider_name_list(self.config)
            combo.clear()
            combo.addItems(names)
            idx = combo.findText(p_name)
            if idx >= 0:
                combo.setCurrentIndex(idx)

    def _browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择作文文件夹")
        if folder:
            self.path_entry.setText(folder)
            self._refresh_queue()

    # ---- Task Queue ----
    def _refresh_queue(self):
        with self.queue_lock:
            self.task_queue = [p for p in self.task_queue if os.path.isdir(p)]
        self._render_queue()

    def _render_queue(self):
        old_rows = {}
        for row in range(self.queue_table.rowCount()):
            path_item = self.queue_table.item(row, 2)
            if path_item:
                old_rows[path_item.text()] = {
                    "step": self.queue_table.item(row, 5).text() if self.queue_table.item(row, 5) else "-",
                    "after_count": self.queue_table.item(row, 6).text() if self.queue_table.item(row, 6) else "-",
                    "status": self.queue_table.item(row, 7).text() if self.queue_table.item(row, 7) else "待完成",
                    "log": self.queue_table.item(row, 8).text() if self.queue_table.item(row, 8) else "等待开始...",
                }
        status_colors = {"待完成": "#cfe2ff", "处理中": "#fff3bf", "已完成": "#d4edda", "失败": "#f8d7da"}

        self.queue_table.setRowCount(0)
        with self.queue_lock:
            task_paths = list(self.task_queue)
        for i, task_path in enumerate(task_paths, start=1):
            folder_name = os.path.basename(task_path)
            student, essay = infer_student_and_essay(folder_name)
            count = count_existing_docx_chars(task_path)
            row = self.queue_table.rowCount()
            self.queue_table.insertRow(row)
            self.queue_table.setItem(row, 0, QTableWidgetItem(str(i)))
            self.queue_table.setItem(row, 1, QTableWidgetItem(student))
            self.queue_table.setItem(row, 2, QTableWidgetItem(task_path))
            self.queue_table.setItem(row, 3, QTableWidgetItem(essay))
            self.queue_table.setItem(row, 4, QTableWidgetItem(count))
            if task_path in old_rows:
                row_state = old_rows[task_path]
                self.queue_table.setItem(row, 5, QTableWidgetItem(row_state["step"]))
                self.queue_table.setItem(row, 6, QTableWidgetItem(row_state["after_count"]))
                self.queue_table.setItem(row, 7, QTableWidgetItem(row_state["status"]))
                self.queue_table.setItem(row, 8, QTableWidgetItem(row_state["log"]))
                bg_color = QColor(status_colors.get(row_state["status"], "#ffffff"))
            elif task_path in self.completed_tasks:
                self.queue_table.setItem(row, 5, QTableWidgetItem("完成"))
                self.queue_table.setItem(row, 6, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 7, QTableWidgetItem("已完成"))
                self.queue_table.setItem(row, 8, QTableWidgetItem("之前已处理完成，跳过"))
                bg_color = QColor("#d4edda")
            else:
                self.queue_table.setItem(row, 5, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 6, QTableWidgetItem("-"))
                self.queue_table.setItem(row, 7, QTableWidgetItem("待完成"))
                self.queue_table.setItem(row, 8, QTableWidgetItem("等待开始..."))
                bg_color = QColor("#cfe2ff")
            for col in range(9):
                item = self.queue_table.item(row, col)
                if item:
                    item.setBackground(bg_color)

    def _add_task(self):
        folder = QFileDialog.getExistingDirectory(self, "选择任务文件夹")
        if not folder:
            return
        folder = os.path.abspath(folder)
        if not os.path.isdir(folder) or not has_images_folder(folder):
            self.log_signal.log_message.emit(f"无效文件夹或无图片：{folder}")
            return
        if folder in self.task_queue:
            self.log_signal.log_message.emit(f"已存在：{folder}")
            return
        with self.queue_lock:
            self.task_queue.append(folder)
            self.finished_tasks.discard(folder)
            self.completed_tasks.discard(folder)
            self.in_progress_tasks.discard(folder)
        self._render_queue()
        self.log_signal.log_message.emit(f"已添加任务：{folder}")

    def _remove_task(self):
        rows = self.queue_table.selectionModel().selectedRows()
        if not rows:
            self.log_signal.log_message.emit("请先选择要删除的队列项")
            return
        for idx in sorted(rows, reverse=True):
            path = self.queue_table.item(idx.row(), 2).text()
            with self.queue_lock:
                if path in self.task_queue:
                    self.task_queue.remove(path)
                # 删除时清除完成标记，重新加入可再处理
                self.completed_tasks.discard(path)
                self.finished_tasks.discard(path)
                self.in_progress_tasks.discard(path)
        self._refresh_queue()

    def _requeue_selected_tasks(self):
        rows = self.queue_table.selectionModel().selectedRows()
        if not rows:
            self.log_signal.log_message.emit("请先选择要重新加入的队列项")
            return

        requeued = 0
        skipped_running = 0
        for idx in rows:
            path_item = self.queue_table.item(idx.row(), 2)
            if not path_item:
                continue
            task_path = path_item.text()
            with self.queue_lock:
                if task_path in self.in_progress_tasks:
                    skipped_running += 1
                    continue
                if task_path not in self.task_queue:
                    self.task_queue.append(task_path)
                self.completed_tasks.discard(task_path)
                self.finished_tasks.discard(task_path)
            self.log_signal.task_status.emit(task_path, "pending", "等待重试", "", "手动重新加入队列")
            requeued += 1

        if requeued:
            self.log_signal.log_message.emit(f"已重新加入 {requeued} 个任务")
        if skipped_running:
            self.log_signal.log_message.emit(f"{skipped_running} 个任务正在处理中，已跳过")
        self._refresh_queue()

    def _load_tasks(self):
        folder = self.path_entry.text().strip()
        if not folder or not os.path.isdir(folder):
            self.log_signal.log_message.emit("当前路径无效")
            return
        candidates = scan_folder_for_tasks(folder)
        added = 0
        with self.queue_lock:
            for p in candidates:
                if p not in self.task_queue:
                    self.task_queue.append(p)
                    self.completed_tasks.discard(p)
                    self.finished_tasks.discard(p)
                    self.in_progress_tasks.discard(p)
                    added += 1
        self.log_signal.log_message.emit(f"已读取并加入 {added} 个任务" if added else "没有新任务")
        self._render_queue()

    def _update_task_status(self, task_path: str, status: str, step: str = "", after_count: str = "", log_msg: str = ""):
        """更新任务状态，支持实时步骤、修改后字数和日志"""
        labels = {"pending": "待完成", "running": "处理中", "done": "已完成", "failed": "失败"}
        colors = {"pending": "#cfe2ff", "running": "#fff3bf", "done": "#d4edda", "failed": "#f8d7da"}
        for row in range(self.queue_table.rowCount()):
            if self.queue_table.item(row, 2) and self.queue_table.item(row, 2).text() == task_path:
                if step:
                    self.queue_table.item(row, 5).setText(step)
                if after_count:
                    self.queue_table.item(row, 6).setText(after_count)
                if status:
                    self.queue_table.item(row, 7).setText(labels.get(status, status))
                if log_msg:
                    old_log = self.queue_table.item(row, 8).text()
                    if old_log == "等待开始...":
                        self.queue_table.item(row, 8).setText(log_msg)
                    else:
                        self.queue_table.item(row, 8).setText(old_log + "\n" + log_msg)
                    self.queue_table.scrollToItem(self.queue_table.item(row, 8))
                if status:
                    for col in range(9):
                        item = self.queue_table.item(row, col)
                        if item:
                            item.setBackground(QColor(colors.get(status, "#ffffff")))
                # 完成的任务加入标记集合
                if status == "done":
                    with self.queue_lock:
                        self.completed_tasks.add(task_path)
                        self.finished_tasks.add(task_path)
                elif status == "failed":
                    with self.queue_lock:
                        self.finished_tasks.add(task_path)
                elif status == "pending":
                    with self.queue_lock:
                        self.completed_tasks.discard(task_path)
                        self.finished_tasks.discard(task_path)
                        self.in_progress_tasks.discard(task_path)
                break

    # ---- Start Processing ----
    def _start_processing(self):
        # 自动读取路径下的任务加入队列
        folder = self.path_entry.text().strip()
        added = 0
        requeued_tasks = []
        if folder and os.path.isdir(folder):
            candidates = scan_folder_for_tasks(folder)
            with self.queue_lock:
                for task_path in candidates:
                    if task_path not in self.task_queue:
                        self.task_queue.append(task_path)
                        self.completed_tasks.discard(task_path)
                        self.finished_tasks.discard(task_path)
                        self.in_progress_tasks.discard(task_path)
                        added += 1
        with self.queue_lock:
            # 每次点击开始时，只跳过已完成任务；失败/未完成任务重新进入本批次。
            for task_path in self.task_queue:
                if task_path not in self.completed_tasks and task_path not in self.in_progress_tasks:
                    if task_path in self.finished_tasks:
                        requeued_tasks.append(task_path)
                    self.finished_tasks.discard(task_path)
            should_start_worker = not self.is_processing
            if should_start_worker:
                self.is_processing = True
        if added:
            self.log_signal.log_message.emit(f"已自动读取 {added} 个任务")
        for task_path in requeued_tasks:
            self.log_signal.task_status.emit(task_path, "pending", "等待重试", "", "重新加入队列")
        self._refresh_queue()

        if not should_start_worker:
            self.log_signal.log_message.emit("当前已有任务在处理，新加入的任务会排队继续处理")
            return

        threading.Thread(target=self._run_processing, daemon=True).start()

    def _run_processing(self):
        cfg = self.config
        cfg.setdefault("OCR", {})
        cfg["OCR"].setdefault("XFYUN", {})
        cfg["OCR"]["PROVIDER"] = "xfyun_handwriting"
        cfg["OCR"]["XFYUN"]["URL"] = self.url_entry.text().strip()
        cfg["OCR"]["XFYUN"]["APPID"] = self.appid_entry.text().strip()
        cfg["OCR"]["XFYUN"]["API_KEY"] = self.apikey_entry.text().strip()
        cfg["OCR"]["XFYUN"].setdefault("LANGUAGE", "cn|en")
        cfg["OCR"]["XFYUN"].setdefault("LOCATION", "false")

        cfg.setdefault("APP", {})
        cfg["APP"]["ROOT_DIR"] = self.path_entry.text().strip()

        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"].setdefault("TASKS", {})

        typo_provider = _ensure_provider_exists(cfg, self.typo_provider_combo.currentText() or "deepseek")
        cfg["LLM"]["PROVIDERS"].setdefault(typo_provider, {})
        cfg["LLM"]["PROVIDERS"][typo_provider]["API_KEY"] = self.typo_api_key.text().strip()
        cfg["LLM"]["PROVIDERS"][typo_provider]["BASE_URL"] = self.typo_base_entry.text().strip()
        cfg["LLM"]["PROVIDERS"][typo_provider].setdefault("MODEL", "deepseek-chat" if typo_provider == "deepseek" else "gpt-4o-mini")

        cfg["LLM"]["TASKS"].setdefault("typo_fix", {})
        cfg["LLM"]["TASKS"]["typo_fix"]["ENABLED"] = self.use_typo_fix.isChecked()
        cfg["LLM"]["TASKS"]["typo_fix"]["PROVIDER"] = typo_provider
        cfg["LLM"]["TASKS"]["typo_fix"]["PROMPT"] = self.typo_prompt_text.toPlainText().strip()

        editor_provider = _ensure_provider_exists(cfg, self.editor_provider_combo.currentText() or "deepseek")
        cfg["LLM"]["PROVIDERS"].setdefault(editor_provider, {})
        cfg["LLM"]["PROVIDERS"][editor_provider]["API_KEY"] = self.editor_api_key.text().strip()
        cfg["LLM"]["PROVIDERS"][editor_provider]["BASE_URL"] = self.editor_base_entry.text().strip()
        cfg["LLM"]["PROVIDERS"][editor_provider].setdefault("MODEL", "deepseek-chat" if editor_provider == "deepseek" else "gpt-4o-mini")

        count_min = count_max = None
        min_text = self.editor_count_min.text().strip()
        max_text = self.editor_count_max.text().strip()
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最小值必须是整数")
                with self.queue_lock:
                    self.is_processing = False
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最大值必须是整数")
                with self.queue_lock:
                    self.is_processing = False
                return
        if count_min is not None and count_max is not None and count_min > count_max:
            self.log_signal.log_message.emit("目标字数范围最小值不能大于最大值")
            with self.queue_lock:
                self.is_processing = False
            return

        cfg["LLM"]["TASKS"].setdefault("editor", {})
        cfg["LLM"]["TASKS"]["editor"]["ENABLED"] = self.use_editor.isChecked()
        cfg["LLM"]["TASKS"]["editor"]["PROVIDER"] = editor_provider
        cfg["LLM"]["TASKS"]["editor"]["PROMPT"] = self.editor_prompt_text.toPlainText().strip()
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max

        cfg["OCR"].setdefault("BAIDU_CORRECTION", {})
        cfg["OCR"]["BAIDU_CORRECTION"]["ENABLED"] = self.use_baidu_correction.isChecked()
        cfg["OCR"]["BAIDU_CORRECTION"]["API_KEY"] = self.baidu_api_key_entry.text().strip()
        cfg["OCR"]["BAIDU_CORRECTION"]["SECRET_KEY"] = self.baidu_secret_key_entry.text().strip()

        if not all([cfg.get("OCR", {}).get("XFYUN", {}).get("URL"), cfg.get("OCR", {}).get("XFYUN", {}).get("APPID"), cfg.get("OCR", {}).get("XFYUN", {}).get("API_KEY"), cfg.get("APP", {}).get("ROOT_DIR")]):
            self.log_signal.log_message.emit("请填写完整的 OCR 配置和文件夹路径")
            with self.queue_lock:
                self.is_processing = False
            return

        if not os.path.isdir(cfg["APP"]["ROOT_DIR"]):
            self.log_signal.log_message.emit("文件夹路径无效")
            with self.queue_lock:
                self.is_processing = False
            return

        save_config(cfg)
        tasks_cfg = cfg.get("LLM", {}).get("TASKS", {})
        if tasks_cfg.get("typo_fix", {}).get("ENABLED"):
            self.log_signal.log_message.emit("AI 错别字修正：已启用")
        if cfg.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED"):
            self.log_signal.log_message.emit("百度图片矫正：已启用")

        def task_status_cb(folder_path, status, step="", log_msg="", after_count=""):
            self.log_signal.task_status.emit(
                folder_path,
                status,
                step or "",
                str(after_count) if after_count else "",
                log_msg or "",
            )

        try:
            self.log_signal.log_message.emit(f"开始处理...（并发数：{self.max_parallel_tasks}）")
            from ocr_main import process_folder

            def process_one_task(task_path):
                task_name = os.path.basename(task_path)

                def task_log(msg):
                    self.log_signal.log_message.emit(f"[{task_name}] {msg}")
                    self.log_signal.task_status.emit(task_path, "", "", "", msg or "")

                try:
                    self.log_signal.task_status.emit(task_path, "running", "排队启动", "", "开始处理")
                    task_log(f"处理: {task_path}")
                    process_folder(
                        task_path,
                        log_callback=task_log,
                        use_typo_fix=bool(tasks_cfg.get("typo_fix", {}).get("ENABLED", False)),
                        use_editor=bool(tasks_cfg.get("editor", {}).get("ENABLED", False)),
                        task_status_callback=task_status_cb,
                    )
                except Exception as exc:
                    self.log_signal.task_status.emit(task_path, "failed", "", "", f"失败: {exc}")
                    self.log_signal.log_message.emit(f"[{task_name}] 处理失败：{exc}")

            started_any = False
            futures = {}
            with ThreadPoolExecutor(max_workers=self.max_parallel_tasks) as executor:
                while True:
                    with self.queue_lock:
                        slots = self.max_parallel_tasks - len(futures)
                        pending_tasks = [
                            p for p in self.task_queue
                            if p not in self.finished_tasks and p not in self.in_progress_tasks
                        ]
                        next_tasks = pending_tasks[:max(0, slots)]
                        for task_path in next_tasks:
                            if os.path.isdir(task_path):
                                self.in_progress_tasks.add(task_path)
                            else:
                                self.finished_tasks.add(task_path)

                    for task_path in next_tasks:
                        if not os.path.isdir(task_path):
                            continue
                        started_any = True
                        futures[executor.submit(process_one_task, task_path)] = task_path

                    if not futures:
                        with self.queue_lock:
                            has_more = any(
                                p not in self.finished_tasks and p not in self.in_progress_tasks
                                for p in self.task_queue
                            )
                            if not has_more:
                                self.is_processing = False
                        if not has_more:
                            if not started_any:
                                self.log_signal.log_message.emit("没有待处理的任务（已完成的任务已跳过）")
                            break

                    done_futures, _ = wait(futures.keys(), timeout=0.3, return_when=FIRST_COMPLETED)
                    for future in done_futures:
                        task_path = futures.pop(future)
                        try:
                            future.result()
                        except Exception as exc:
                            self.log_signal.task_status.emit(task_path, "failed", "", "", f"失败: {exc}")
                            self.log_signal.log_message.emit(f"[{os.path.basename(task_path)}] 处理失败：{exc}")
                        finally:
                            with self.queue_lock:
                                self.in_progress_tasks.discard(task_path)
                                self.finished_tasks.add(task_path)

            self.log_signal.log_message.emit("全部处理完成")
        except Exception as e:
            self.log_signal.log_message.emit(f"处理失败：{e}")
        finally:
            with self.queue_lock:
                self.in_progress_tasks.clear()
                self.is_processing = False

    # ===================== PAGE 2: AI DOCX =====================
    def _init_page_ai(self):
        layout = QVBoxLayout(self.page_ai)
        layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(4, 4, 4, 4)
        scroll_layout.setSpacing(6)

        # AI Config
        ai_group = QGroupBox("AI 配置")
        ai_layout = QVBoxLayout()

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("处理文件夹"))
        self.ai_path_entry = QLineEdit()
        row1.addWidget(self.ai_path_entry, 1)
        btn_browse = QPushButton("浏览")
        btn_browse.setFixedWidth(70)
        btn_browse.clicked.connect(self._browse_ai_folder)
        row1.addWidget(btn_browse)
        ai_layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.addWidget(QLabel("AI Provider"))
        self.ai_provider_combo = QComboBox()
        self.ai_provider_combo.addItems(_provider_name_list(self.config))
        ai_provider = _normalize_provider_name((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")) or "deepseek"
        idx = self.ai_provider_combo.findText(ai_provider)
        if idx >= 0:
            self.ai_provider_combo.setCurrentIndex(idx)
        self.ai_provider_combo.currentTextChanged.connect(self._on_ai_provider_change)
        row2.addWidget(self.ai_provider_combo)
        btn_new_ai = QPushButton("+ 新增")
        btn_new_ai.setFixedWidth(68)
        btn_new_ai.clicked.connect(lambda: self._add_provider(self.ai_provider_combo))
        row2.addWidget(btn_new_ai)
        ai_layout.addLayout(row2)

        row3 = QHBoxLayout()
        row3.addWidget(QLabel("AI API Key"))
        self.ai_key_entry = QLineEdit()
        self.ai_key_entry.setEchoMode(QLineEdit.Password)
        self.ai_key_entry.setText((self.config.get("LLM", {}).get("PROVIDERS", {}).get(ai_provider, {}) or {}).get("API_KEY", ""))
        row3.addWidget(self.ai_key_entry, 1)
        ai_layout.addLayout(row3)

        row4 = QHBoxLayout()
        row4.addWidget(QLabel("Base URL"))
        self.ai_url_entry = QLineEdit((self.config.get("LLM", {}).get("PROVIDERS", {}).get(ai_provider, {}) or {}).get("BASE_URL", ""))
        row4.addWidget(self.ai_url_entry, 1)
        ai_layout.addLayout(row4)

        row5 = QHBoxLayout()
        row5.addWidget(QLabel("AI Prompt"))
        self.ai_prompt_text = QTextEdit()
        self.ai_prompt_text.setMaximumHeight(80)
        self.ai_prompt_text.setPlainText((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"))
        row5.addWidget(self.ai_prompt_text, 1)
        ai_layout.addLayout(row5)

        row6 = QHBoxLayout()
        row6.addWidget(QLabel("目标字数"))
        self.ai_count_min = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""))
        self.ai_count_min.setFixedWidth(100)
        row6.addWidget(self.ai_count_min)
        row6.addWidget(QLabel("-"))
        self.ai_count_max = QLineEdit(str((self.config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""))
        self.ai_count_max.setFixedWidth(100)
        row6.addWidget(self.ai_count_max)
        row6.addWidget(QLabel("（空白表示自动）"))
        row6.addStretch()
        ai_layout.addLayout(row6)

        ai_group.setLayout(ai_layout)
        scroll_layout.addWidget(ai_group)

        # Task flow
        flow_group = QGroupBox("处理流程（勾选/取消步骤，拖动上下移动顺序）")
        flow_layout = QVBoxLayout()

        self.task_config = [
            {"id": "6", "name": "6. 转换 DOC -> DOCX", "enabled": True, "order": 0},
            {"id": "1", "name": "1. 清除空格", "enabled": True, "order": 1},
            {"id": "AI", "name": "AI 改作文", "enabled": True, "order": 2},
            {"id": "2", "name": '2. 添加"修改前/后"', "enabled": True, "order": 3},
            {"id": "3", "name": "3. 格式化字体段落", "enabled": True, "order": 4},
            {"id": "5", "name": "5. 修改作者", "enabled": True, "order": 5},
        ]

        self.task_checkboxes = {}
        for task in sorted(self.task_config, key=lambda x: x["order"]):
            cb = QCheckBox(task["name"])
            cb.setChecked(task["enabled"])
            cb.toggled.connect(lambda checked, t=task: t.update({"enabled": checked}))
            flow_layout.addWidget(cb)
            self.task_checkboxes[task["id"]] = cb

        flow_group.setLayout(flow_layout)
        scroll_layout.addWidget(flow_group)

        # Start button
        btn_start_ai = QPushButton("🚀 开始流程")
        btn_start_ai.setStyleSheet("background-color: #2196F3; color: white; font-size: 14px; padding: 8px;")
        btn_start_ai.clicked.connect(self._start_ai_workflow)
        scroll_layout.addWidget(btn_start_ai)

        # Log (collapsible)
        self.ai_log_section = CollapsibleSection("处理日志", collapsed=True)
        self.ai_log_text = QTextEdit()
        self.ai_log_text.setReadOnly(True)
        self.ai_log_text.setMaximumHeight(200)
        self.ai_log_section.add_widget(self.ai_log_text)
        scroll_layout.addWidget(self.ai_log_section)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)

    def _browse_ai_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择处理文件夹")
        if folder:
            self.ai_path_entry.setText(folder)

    def _on_ai_provider_change(self, name):
        p_name = _ensure_provider_exists(self.config, name)
        provider_cfg = (self.config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {})
        self.ai_key_entry.setText(provider_cfg.get("API_KEY", ""))
        self.ai_url_entry.setText(provider_cfg.get("BASE_URL", ""))

    def _start_ai_workflow(self):
        threading.Thread(target=self._run_ai_workflow, daemon=True).start()

    def _run_ai_workflow(self):
        folder = self.ai_path_entry.text().strip()
        selected_provider = _ensure_provider_exists(self.config, self.ai_provider_combo.currentText() or "deepseek")
        api_key = self.ai_key_entry.text().strip()
        base_url = self.ai_url_entry.text().strip()
        prompt = self.ai_prompt_text.toPlainText().strip() or None

        min_text = self.ai_count_min.text().strip()
        max_text = self.ai_count_max.text().strip()
        count_min = count_max = None
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最小值必须是整数")
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                self.log_signal.log_message.emit("目标字数范围最大值必须是整数")
                return

        if not folder or not api_key:
            self._append_log_ai("请填写文件夹路径和 API Key")
            return
        if not os.path.isdir(folder):
            self._append_log_ai("文件夹路径无效")
            return

        cfg = self.config
        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"].setdefault("TASKS", {})
        cfg["LLM"]["PROVIDERS"].setdefault(selected_provider, {})
        cfg["LLM"]["PROVIDERS"][selected_provider]["API_KEY"] = api_key
        cfg["LLM"]["PROVIDERS"][selected_provider]["BASE_URL"] = base_url
        cfg["LLM"]["PROVIDERS"][selected_provider].setdefault("MODEL", "deepseek-chat" if selected_provider == "deepseek" else "gpt-4o-mini")
        cfg["LLM"]["TASKS"].setdefault("editor", {})
        cfg["LLM"]["TASKS"]["editor"]["PROMPT"] = prompt or "{text}"
        cfg["LLM"]["TASKS"]["editor"]["ENABLED"] = True
        cfg["LLM"]["TASKS"]["editor"]["PROVIDER"] = selected_provider
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        cfg["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max
        save_config(cfg)

        self._append_log_ai("开始处理流程...")
        self._append_log_ai("【准备】复制原始文件...")
        import shutil
        copied_files = []
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif")
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if (name_lower.endswith(".docx") or name_lower.endswith(image_exts)) and not name_check.startswith("~$") and not name_check.startswith("改 "):
                    original_path = os.path.join(root, file)
                    new_filename = f"改 {file}"
                    new_path = os.path.join(root, new_filename)
                    try:
                        shutil.copy2(original_path, new_path)
                        copied_files.append(new_filename)
                        self._append_log_ai(f"  {new_filename}")
                    except Exception as e:
                        self._append_log_ai(f"  {file} 复制失败: {e}")

        if not copied_files:
            self._append_log_ai("未找到需要处理的文件")
            return

        enabled_tasks = sorted([(t["id"], t["order"]) for t in self.task_config if t["enabled"]], key=lambda x: x[1])
        task_step_names = {"6": "DOC转DOCX", "1": "清除空格", "AI": "AI改作文", "2": "添加标签", "3": "格式化", "5": "改作者"}

        def update_step(task_path, step_name):
            QTimer.singleShot(0, lambda p=task_path, s=step_name: self._update_task_status(p, "running", step=step_name))

        try:
            for task_id, _ in enabled_tasks:
                step_name = task_step_names.get(task_id, task_id)
                self._append_log_ai(f"【{step_name}】开始...")
                # 更新所有任务的当前步骤
                for tp in copied_files:
                    folder_of_file = os.path.join(folder, os.path.dirname(tp))
                    update_step(folder_of_file, step_name)

                if task_id == "6":
                    self._convert_docs(folder)
                elif task_id == "1":
                    self._clear_spaces(folder)
                elif task_id == "AI":
                    self._process_ai(folder, api_key, base_url, prompt, count_min=count_min, count_max=count_max)
                elif task_id == "2":
                    self._add_labels(folder)
                elif task_id == "3":
                    self._format_docs(folder)
                elif task_id == "5":
                    self._set_author(folder)

            # 更新完成状态和修改后字数
            for root_dir, files in iter_files_limited(folder, max_depth=4):
                for file in files:
                    if file.startswith("改 ") and file.endswith(".docx"):
                        full_path = os.path.join(root_dir, file)
                        after_count = count_existing_docx_chars(os.path.dirname(full_path))
                        QTimer.singleShot(0, lambda p=full_path, c=after_count: self._update_task_status(p, "done", step="完成", after_count=c))

            self._append_log_ai("所有流程完成！")
        except Exception as e:
            self._append_log_ai(f"处理失败：{e}")
            import traceback
            traceback.print_exc()

    def _convert_docs(self, folder):
        import subprocess
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                name_lower = file.lower()
                if name_lower.endswith(".doc") and not file.startswith("~$"):
                    doc_path = os.path.join(root, file)
                    try:
                        subprocess.run(["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", root], capture_output=True, timeout=30)
                        base_name = os.path.splitext(os.path.basename(doc_path))[0]
                        new_path = os.path.join(root, base_name + ".docx")
                        if os.path.exists(new_path):
                            os.remove(doc_path)
                            self._append_log_ai(f"  {base_name}")
                    except Exception as e:
                        self._append_log_ai(f"  {file}: {e}")

    def _clear_spaces(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if file.lower().endswith(".docx") and not file.startswith("~$") and not file.startswith("改 "):
                    continue
                if file.lower().endswith(".docx") and not file.startswith("~$"):
                    try:
                        doc = Document(os.path.join(root, file))
                        for para in doc.paragraphs:
                            for run in para.runs:
                                run.text = run.text.strip()
                        doc.save(os.path.join(root, file))
                        self._append_log_ai(f"  {file}")
                    except Exception as e:
                        self._append_log_ai(f"  {file}: {e}")

    def _process_ai(self, folder, api_key, base_url, prompt_template, count_min=None, count_max=None):
        if not prompt_template:
            prompt_template = "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"

        client = OpenAI(api_key=api_key, base_url=base_url)
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$"):
                    continue
                if not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() and p.text.strip() not in ("修改前：", "修改后：")])
                    if not all_text.strip():
                        self._append_log_ai(f"  {file} (空文档)")
                        continue

                    original_count = count_chinese_characters(all_text)
                    if count_min is None or count_max is None:
                        default_min, default_max = determine_word_count_bounds(original_count)
                        count_min = count_min if count_min is not None else default_min
                        count_max = count_max if count_max is not None else default_max

                    for attempt in range(1, 5):
                        if "{text}" in prompt_template:
                            current_prompt = prompt_template.format(text=all_text)
                        else:
                            current_prompt = prompt_template + "\n\n" + all_text
                        current_prompt += f"\n\n请注意：这一次的修改后的正文总字数应控制在 {count_min} 到 {count_max} 之间"
                        if attempt > 1:
                            current_prompt += f"\n\n字数不符合规则，请重新修改并返回修改后的正文。只输出正文，不要解释。这次要求字数在 {count_min} 到 {count_max} 之间。"

                        self._append_log_ai(f"  {file} AI 第{attempt}次输出，正在检查字数...")
                        response = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "system", "content": "你是一名严谨的中文校对助手"}, {"role": "user", "content": current_prompt}], temperature=0.1, stream=False)
                        result_text = response.choices[0].message.content.strip()
                        current_count = count_chinese_characters(result_text)
                        if count_min <= current_count <= count_max:
                            self._append_log_ai(f"  {file} 字数符合：{current_count}（目标 {count_min}-{count_max}）")
                            break
                        self._append_log_ai(f"  {file} 字数不合规：{current_count}，目标 {count_min}-{count_max}，正在重试...")
                    else:
                        raise RuntimeError(f"{file} AI 输出字数不符合要求")

                    last_para = doc.paragraphs[-1] if doc.paragraphs else None
                    if last_para:
                        if last_para.runs:
                            last_para.runs[-1].add_break(WD_BREAK.PAGE)
                        else:
                            last_para.add_run().add_break(WD_BREAK.PAGE)
                    para_modify = doc.add_paragraph("修改后：")
                    para_modify.paragraph_format.first_line_indent = Cm(0.74)
                    para_modify.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    para_modify.paragraph_format.line_spacing = Pt(12)
                    for line in result_text.split("\n"):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Cm(0.74)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            p.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self._append_log_ai(f"  {file}")
                except Exception as e:
                    self._append_log_ai(f"  {file}: {e}")

    def _add_labels(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    if doc.paragraphs:
                        last_para = doc.paragraphs[-1]
                        has_modify = last_para.text.strip() == "修改后：" or (len(doc.paragraphs) > 1 and doc.paragraphs[-2].text.strip() == "修改后：")
                        if doc.paragraphs[0].text.strip() != "修改前：":
                            doc.paragraphs[0].insert_paragraph_before("修改前：")
                        if not has_modify:
                            last_para = doc.paragraphs[-1]
                            if last_para.runs:
                                last_para.runs[-1].add_break(WD_BREAK.PAGE)
                            else:
                                last_para.add_run().add_break(WD_BREAK.PAGE)
                            para = doc.add_paragraph("修改后：")
                            para.paragraph_format.first_line_indent = Cm(0.74)
                            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self._append_log_ai(f"  {file}")
                except Exception as e:
                    self._append_log_ai(f"  {file}: {e}")

    def _format_docs(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    style = doc.styles["Normal"]
                    style.font.name = "宋体"
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    style.font.size = Pt(12)
                    for para in doc.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0.74)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    self._append_log_ai(f"  {file}")
                except Exception as e:
                    self._append_log_ai(f"  {file}: {e}")

    def _set_author(self, folder):
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                if not file.lower().endswith(".docx") or file.startswith("~$") or not file.startswith("改 "):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    doc.core_properties.author = "思睿教育_美丽可爱的尹老师"
                    doc.save(doc_path)
                    self._append_log_ai(f"  {file}")
                except Exception as e:
                    self._append_log_ai(f"  {file}: {e}")


# ===================== Main =====================
if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
