import os
import re
import subprocess
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from openai import OpenAI

# ========== 清理文本 ==========
def clean_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    text = text.replace('\u3000', '')
    return text

# ========== 只清除空格 ==========
def clear_spaces(doc):
    for para in doc.paragraphs:
        cleaned_text = clean_text(para.text)
        if para.text != cleaned_text:
            para.text = cleaned_text

# ========== 添加“修改前/修改后”标签 ==========
def add_modify_labels(doc):
    doc.paragraphs[0].insert_paragraph_before("修改前：")
    last_para = doc.paragraphs[-1]
    if last_para.runs:
        last_para.runs[-1].add_break(WD_BREAK.PAGE)
    else:
        last_para.add_run().add_break(WD_BREAK.PAGE)
    para_modify_after = doc.add_paragraph("修改后：")
    para_modify_after.paragraph_format.first_line_indent = Cm(0.74)
    para_modify_after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    para_modify_after.paragraph_format.line_spacing = Pt(12)
    para_modify_after.paragraph_format.space_before = Pt(0)
    para_modify_after.paragraph_format.space_after = Pt(0)

# ========== 格式化字体段落 ==========
def format_style(doc):
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    for para in doc.paragraphs:
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para.paragraph_format.line_spacing = Pt(12)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

# ========== 修改作者 ==========
def set_author(doc, author="思睿教育_美丽可爱的尹老师"):
    doc.core_properties.author = author

# ========== 格式化 DOCX ==========
def format_docx(doc_path, clear_space=False, add_labels=False, style_format=False, change_author=False):
    try:
        doc = Document(doc_path)
        if clear_space:
            clear_spaces(doc)
        if add_labels:
            add_modify_labels(doc)
        if style_format:
            format_style(doc)
        if change_author:
            set_author(doc)
        doc.save(doc_path)
        print(f"✔ 已格式化：{doc_path}")
    except Exception as e:
        print(f"⚠ 格式化失败：{doc_path}，错误：{e}")

# ========== LibreOffice 转换 DOC→DOCX ==========
def convert_doc_to_docx_libreoffice(doc_path, output_folder):
    try:
        os.makedirs(output_folder, exist_ok=True)
        cmd = ['soffice', '--headless', '--convert-to', 'docxx', doc_path, '--outdir', output_folder]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode == 0:
            base_name = os.path.splitext(os.path.basename(doc_path))[0]
            new_path = os.path.join(output_folder, base_name + ".docxx")
            if os.path.exists(new_path):
                print(f"✔ 转换成功：{doc_path} → {new_path}")
                return new_path
            else:
                print(f"⚠ 转换后文件未找到：{new_path}")
                return None
        else:
            print(f"⚠ 转换失败：{doc_path}\n{result.stderr.strip()}")
            return None
    except Exception as e:
        print(f"⚠ 转换异常：{doc_path}，错误：{e}")
        return None

# ========== 处理文件夹 ==========
def process_folder(input_folder, output_folder, 
                   do_clear_space=False, do_add_labels=False, do_style_format=False, do_change_author=False, do_convert_only=False):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    log_success = []
    log_fail = []
    for root, _, files in os.walk(input_folder):
        for file in files:
            full_path = os.path.join(root, file)
            try:
                if file.lower().endswith(".doc") and not file.startswith("~$"):
                    if do_convert_only:
                        new_path = convert_doc_to_docx_libreoffice(full_path, root)  # 改：转换后就在原文件夹
                        if new_path:
                            log_success.append(full_path)
                            try:
                                os.remove(full_path)
                                print(f"🗑 已删除原文件：{full_path}")
                            except Exception as e:
                                print(f"⚠ 删除失败：{full_path}，错误：{e}")
                        else:
                            log_fail.append(full_path)
                    else:
                        new_path = convert_doc_to_docx_libreoffice(full_path, output_folder)
                        if new_path:
                            format_docx(new_path, do_clear_space, do_add_labels, do_style_format, do_change_author)
                            log_success.append(full_path)
                            try:
                                os.remove(full_path)
                                print(f"🗑 已删除原文件：{full_path}")
                            except Exception as e:
                                print(f"⚠ 删除失败：{full_path}，错误：{e}")
                        else:
                            log_fail.append(full_path)
                elif file.lower().endswith(".docxx") and not file.startswith("~$"):
                    if do_convert_only:
                        continue
                    format_docx(full_path, do_clear_space, do_add_labels, do_style_format, do_change_author)
                    log_success.append(full_path)
            except Exception as e:
                print(f"⚠ 处理失败：{full_path}，错误：{e}")
                log_fail.append(full_path)
    with open("处理日志.txt", "w", encoding="utf-8") as log:
        log.write("成功处理文件：\n")
        log.write("\n".join(log_success))
        log.write("\n\n失败文件：\n")
        log.write("\n".join(log_fail))
    print("\n✅ 所有任务完成！日志保存在：处理日志.txt")
# ========== 功能 8：AI 处理流程（6→1→3→AI→2→3→5）==========
def process_with_ai(input_folder, api_key: str, base_url: str = None, prompt_template: str = None):
    """
    综合 AI 处理流程：
    1. 转换 DOC→DOCX (功能 6)
    2. 清除空格 (功能 1)
    3. 格式化 (功能 3)
    4. 读取所有 DOCX 内容 → 发送给 AI → 获取结果
    5. 将 AI 结果写回文档
    6. 添加标签 (功能 2)
    7. 格式化 (功能 3)
    8. 修改作者 (功能 5)
    """
    if not api_key:
        print("❌ 未提供 API Key，无法进行 AI 处理")
        return
    
    client = OpenAI(api_key=api_key, base_url=(base_url or "https://api.deepseek.com"))
    
    if not prompt_template:
        prompt_template = (
            "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n"
            "要求：\n"
            "1. 不改变原意\n"
            "2. 不润色文风\n"
            "3. 不增删内容\n"
            "4. 保持原有段落结构\n"
            "5. 只输出修改后的完整文章正文\n"
            "\n{text}"
        )
    
    print(f"\n📋 开始 AI 处理流程...")
    
    # ===== 第 1 阶段：转换 + 清除空格 + 初步格式化 =====
    print("\n【第 1 阶段】转换 DOC→DOCX、清除空格、初步格式化...")
    process_folder(input_folder, input_folder, 
                   do_convert_only=False,
                   do_clear_space=True, 
                   do_style_format=True)
    
    # ===== 第 2 阶段：处理所有 DOCX 文件 =====
    print("\n【第 2 阶段】读取文档内容并发送给 AI...")
    docx_files = []
    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.lower().endswith(".docx") and not file.startswith("~$"):
                docx_files.append(os.path.join(root, file))
    
    if not docx_files:
        print("❌ 未找到任何 .docx 文件")
        return
    
    for doc_path in docx_files:
        print(f"\n处理: {doc_path}")
        try:
            doc = Document(doc_path)
            
            # 读取所有段落内容
            all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            if not all_text.strip():
                print(f"  ⚠️ 文档为空，跳过")
                continue
            
            print(f"  📤 发送给 AI （共 {len(all_text)} 字符）...")
            
            # 格式化 prompt
            if "{text}" in prompt_template:
                full_prompt = prompt_template.format(text=all_text)
            else:
                full_prompt = prompt_template + "\n\n" + all_text
            
            # 调用 AI
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一名严谨的中文校对助手"},
                    {"role": "user", "content": full_prompt},
                ],
                temperature=0.1,
                stream=False
            )
            
            ai_result = response.choices[0].message.content.strip()
            print(f"  ✅ 收到 AI 返回 （共 {len(ai_result)} 字符）")
            
            # ===== 将 AI 结果写入文档 =====
            # 先清空现有段落
            for para in list(doc.paragraphs):
                if para.text.strip():
                    p = para._element
                    p.getparent().remove(p)
            
            # 插入 AI 返回的段落
            for line in ai_result.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    fmt = p.paragraph_format
                    fmt.first_line_indent = Cm(0.74)
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    fmt.line_spacing = Pt(12)
            
            doc.save(doc_path)
            print(f"  💾 已保存 AI 结果到文档")
            
        except Exception as e:
            print(f"  ❌ 处理失败: {e}")
            import traceback
            traceback.print_exc()
    
    # ===== 第 3 阶段：添加标签、格式化、修改作者 =====
    print("\n【第 3 阶段】添加标签、格式化、修改作者...")
    process_folder(input_folder, input_folder,
                   do_add_labels=True,
                   do_style_format=True,
                   do_change_author=True)
    
    print("\n✅ AI 处理流程全部完成！")
# ========== 功能 7：合并“修改后”段落 ==========
def merge_modified_paragraphs_with_format(input_folder, output_file):
    combined_doc = Document()
    count = 0
    pattern = re.compile(r"修改后[:：]\s*", re.IGNORECASE)
    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.lower().endswith(".docxx") and not file.startswith("~$"):
                file_path = os.path.join(root, file)
                try:
                    doc = Document(file_path)
                except Exception as e:
                    print(f"⚠ 无法读取 {file_path}，跳过：{e}")
                    continue
                start_index = None
                for i, p in enumerate(doc.paragraphs):
                    if pattern.search(p.text.strip()):
                        start_index = i
                        break
                if start_index is not None:
                    count += 1
                    for p in doc.paragraphs[start_index:]:
                        text = p.text.strip()
                        if not text:
                            continue
                        new_para = combined_doc.add_paragraph()
                        for run in p.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                    combined_doc.add_page_break()
    # 格式化合并结果
    style = combined_doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    for para in combined_doc.paragraphs:
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        para.paragraph_format.line_spacing = Pt(12)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
    if count > 0:
        combined_doc.save(output_file)
        print(f"🎉 合并完成！共收集到 {count} 篇文章，已保存到：{output_file}")
    else:
        print("❌ 没有找到任何包含“修改后：”或“修改后:”的文档。")

# ========== 主程序 ==========
def main():
    input_dir = input("请输入源文件夹路径：").strip('" ')
    output_dir = input_dir
    if not os.path.isdir(input_dir):
        print("❌ 输入路径无效，请检查。")
        return
    while True:
        print("\n请选择操作：")
        print("1. 只清除*.docx中空格")
        print('2. 只在*.docx添加"修改前："和"修改后："')
        print("3. 只格式化*.docx字体字号段落")
        print("4. 全部运行（转换+清除空格+添加标签+格式化+修改作者）")
        print("5. 只修改*.docx文件作者")
        print("6. 只转换 DOC → DOCX (转换后放回原目录)")
        print('7. 收集所有包含"修改后"段落的文章到一个文件')
        print("8. AI 综合处理（6→1→3→AI修正→2→3→5）")
        print("0. 退出")
        choice = input("请输入数字选择：").strip()
        if choice == '1':
            process_folder(input_dir, output_dir, do_clear_space=True)
        elif choice == '2':
            process_folder(input_dir, output_dir, do_add_labels=True)
        elif choice == '3':
            process_folder(input_dir, output_dir, do_style_format=True)
        elif choice == '4':
            process_folder(input_dir, output_dir, 
                           do_clear_space=True, do_add_labels=True, do_style_format=True, do_change_author=True)
        elif choice == '5':
            process_folder(input_dir, output_dir, do_change_author=True)
        elif choice == '6':
            process_folder(input_dir, output_dir, do_convert_only=True)
        elif choice == '7':
            output_file = os.path.join(input_dir, "合并后文档.docxx")
            merge_modified_paragraphs_with_format(input_dir, output_file)
        elif choice == '8':
            api_key = input("请输入 AI API Key: ").strip()
            base_url = input("请输入 API Base URL（默认 DeepSeek）: ").strip()
            base_url = base_url or "https://api.deepseek.com"
            print("\n【可选】输入自定义 Prompt（留空使用默认）:")
            prompt = input().strip()
            process_with_ai(input_dir, api_key, base_url=base_url, prompt_template=prompt if prompt else None)
        elif choice == '0':
            print("退出程序。")
            break
        else:
            print("无效选择，请重新输入。")

if __name__ == "__main__":
    main()
