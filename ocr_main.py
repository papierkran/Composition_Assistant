# coding=utf-8
# ocr_main.py
import os
import time
import hashlib
import base64
import argparse
import logging
from pathlib import Path

import requests
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from openai import OpenAI

import json

from config_migrate import ensure_new_schema
from llm_client import resolve_task_client, resolve_task_clients

# 默认使用项目目录下的 config.json；如需覆盖，可通过环境变量 OCR_CONFIG_FILE 指定。
CONFIG_FILE = Path(os.environ.get("OCR_CONFIG_FILE", "config.json")).expanduser()

def load_config(path: Path = None):
    """Load configuration from JSON file.

    Tries environment-configured path first, then falls back to './config.json'.
    Returns DEFAULT-like dict when file missing will raise a RuntimeError to force user action.
    """
    cfg_path = Path(path or CONFIG_FILE)
    if not cfg_path.exists():
        # try local config.json as a fallback
        local = Path("config.json")
        if local.exists():
            cfg_path = local
        else:
            raise RuntimeError("❌ 未找到 config.json 或相关配置文件，请先配置")

    with cfg_path.open("r", encoding="utf-8") as f:
        return json.load(f)

config = ensure_new_schema(load_config())

# global debug flag driven from config (APP.DEBUG) or env
DEBUG = bool(config.get("APP", {}).get("DEBUG", False)) or os.environ.get("OCR_DEBUG", "") != ""
_LOGGER = logging.getLogger("ocr_main")
if DEBUG:
    logging.basicConfig(level=logging.DEBUG)
else:
    logging.basicConfig(level=logging.INFO)



# ========== OCR（讯飞）配置区 =================================================
OCR_CONFIG = (config.get("OCR", {}) or {}).get("XFYUN", {})

URL = OCR_CONFIG.get("URL")
APPID = OCR_CONFIG.get("APPID")
API_KEY = OCR_CONFIG.get("API_KEY")
language = OCR_CONFIG.get("LANGUAGE", "cn|en")
location = OCR_CONFIG.get("LOCATION", "false")

if not all([URL, APPID, API_KEY]):
    raise RuntimeError("❌ OCR 配置不完整，请检查 config.json")

# ==================================================================


# ========== LLM（OpenAI-compatible）任务配置 =================================================
# 兼容：通过 ensure_new_schema 已把旧配置映射到 LLM.PROVIDERS + LLM.TASKS
DEFAULT_TYPO_FIX_PROMPT = (
    "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n"
    "要求：\n"
    "1. 不改变原意\n"
    "2. 不润色文风\n"
    "3. 不增删内容\n"
    "4. 保持原有段落结构\n"
    "5. 只输出修改后的完整文章正文\n"
    "6. 格式应该是  标题  （\\n）下一行  ——xx(替换为姓名)  然后文章内容\n"
    "标题不要出现 ‘题目：’ ‘标题：’等字样\n\n"
    "{text}"
)
# ==================================================================




# 讯飞ocr相关设置
def getHeader() -> dict:
    """Build request header for 讯飞 OCR API.

    Uses global `API_KEY` and `APPID` from configuration.
    """
    curTime = str(int(time.time()))
    param = json.dumps({"language": language, "location": location})
    paramBase64 = base64.b64encode(param.encode('utf-8')).decode('utf-8')
    checkSum_str = API_KEY + curTime + paramBase64
    checkSum = hashlib.md5(checkSum_str.encode('utf-8')).hexdigest()
    header = {
        'X-CurTime': curTime,
        'X-Param': paramBase64,
        'X-Appid': APPID,
        'X-CheckSum': checkSum,
        'Content-Type': 'application/x-www-form-urlencoded; charset=utf-8',
    }
    return header


def _update_globals_from_config(cfg: dict):
    """Update module-level globals from a configuration dict.

    This is used by the CLI `main()` to reload settings when a different
    config file is provided at runtime.
    """
    global config, OCR_CONFIG, URL, APPID, API_KEY, language, location

    config = ensure_new_schema(cfg)
    OCR_CONFIG = (config.get("OCR", {}) or {}).get("XFYUN", {})
    URL = OCR_CONFIG.get("URL")
    APPID = OCR_CONFIG.get("APPID")
    API_KEY = OCR_CONFIG.get("API_KEY")
    language = OCR_CONFIG.get("LANGUAGE", "cn|en")
    location = OCR_CONFIG.get("LOCATION", "false")


def main(argv=None):
    """Command-line entry point.

    Supports overriding config path, root directory, toggling AI steps and debug.
    """
    parser = argparse.ArgumentParser(description="Composition OCR Assistant - CLI")
    parser.add_argument("root", nargs="?", help="要处理的根目录路径（可选，优先）")
    parser.add_argument("--config", "-c", help="配置文件路径（JSON）")
    parser.add_argument("--no-deepseek", action="store_true", help="禁用 DeepSeek 调用")
    parser.add_argument("--no-editor", action="store_true", help="禁用第二步编辑 API")
    parser.add_argument("--debug", action="store_true", help="启用调试日志")

    args = parser.parse_args(argv)

    # reload config if provided
    if args.config:
        cfg = load_config(Path(args.config))
        _update_globals_from_config(cfg)

    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)

    root_dir = args.root or config.get("APP", {}).get("ROOT_DIR")
    if not root_dir:
        raise RuntimeError("未指定要处理的目录。请提供命令行参数或在配置中设置 APP.ROOT_DIR")

    # Legacy CLI flags are kept; they map to new LLM task toggles.
    llm_tasks = (config.get("LLM", {}) or {}).get("TASKS", {})
    typo_enabled = bool(llm_tasks.get("typo_fix", {}).get("ENABLED", False)) and (not args.no_deepseek)
    editor_enabled = bool(llm_tasks.get("editor", {}).get("ENABLED", False)) and (not args.no_editor)

    process_all(
        root_dir,
        log_callback=print,
        use_typo_fix=typo_enabled,
        use_editor=editor_enabled,
    )



# LLM（OpenAI-compatible）相关设置

def llm_fix_typos(prompt_template: str, text: str, client: OpenAI, model: str) -> str:
    """调用 OpenAI-compatible LLM，只修改错别字，不改变原意，不润色。

    `prompt_template` 可以包含 `{text}` 占位符；如果没有，会把 `text` 追加到末尾。
    是否打印 prompt/response 由 DEBUG 控制。
    """
    if not prompt_template:
        prompt_template = DEFAULT_TYPO_FIX_PROMPT

    if "{text}" in prompt_template:
        prompt = prompt_template.format(text=text)
    else:
        prompt = prompt_template + "\n\n" + text

    if DEBUG:
        print("\n=================【发送给 LLM 的内容】=================\n")
        print(prompt)
        print("\n======================================================\n")

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "你是一名严谨的中文校对助手"},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        stream=False,
    )

    if DEBUG:
        print("\n=================【LLM 原始返回 response】=================\n")
        try:
            print(response)
        except Exception:
            print(repr(response))
        print("\n=======================================================\n")

    result_text = response.choices[0].message.content.strip()

    if DEBUG:
        print("\n=================【LLM 修改后的正文】=================\n")
        print(result_text)
        print("\n=====================================================\n")

    return result_text


# ocr 并处理自然段
def ocr_and_extract_text(image_path: str, timeout: int = 30):
    with open(image_path, 'rb') as f:
        imgfile = f.read()

    data = {'image': base64.b64encode(imgfile).decode('utf-8')}
    headers = getHeader()
    try:
        resp = requests.post(URL, headers=headers, data=data, timeout=timeout)
        resp.raise_for_status()
        result = resp.json()
    except requests.RequestException as e:
        _LOGGER.exception("OCR request failed")
        raise RuntimeError(f"OCR 请求失败: {e}") from e
    # print("OCR 原始返回：", result)

    if result.get('code') != '0':
        raise RuntimeError(f"OCR 失败: {result.get('desc')}")

    # ========== ① 提取所有行 ==========
    lines = []
    for block in result.get("data", {}).get("block", []):
        if block.get("type") != "text":
            continue
        for line in block.get("line", []):
            text = "".join(w.get("content", "") for w in line.get("word", []))
            if text.strip():
                lines.append(text)

    # ========== ② 智能合并自然段 ==========
    paragraphs = []
    buffer = ""

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if len(line) <= 4:
            buffer += line
            continue

        buffer += line

        if buffer[-1] in "。！？…" and len(buffer) >= 30:
            paragraphs.append(buffer)
            buffer = ""

    if buffer:
        paragraphs.append(buffer)

    return paragraphs   # ✅ 关键：返回段落列表


#文档处理配置
def create_word(doc_path: str, all_text_blocks, folder_display_name: str):

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # 修改前
    para_before = doc.add_paragraph("修改前：")
    para_before.paragraph_format.first_line_indent = Cm(0.74)
    para_before.paragraph_format.space_before = Pt(0)
    para_before.paragraph_format.space_after = Pt(0)
    para_before.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    para_before.paragraph_format.line_spacing = Pt(12)

    # 姓名
    para_name = doc.add_paragraph(f"——{folder_display_name}")
    para_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_name.paragraph_format.space_before = Pt(0)
    para_name.paragraph_format.space_after = Pt(0)
    para_name.paragraph_format.line_spacing = Pt(12)

    for para_text in all_text_blocks:
        # OCR 段落
        p = doc.add_paragraph(para_text)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        p.paragraph_format.line_spacing = Pt(12)

    doc.add_page_break()
    # 修改后
    para_after = doc.add_paragraph("修改后：")
    para_after.paragraph_format.first_line_indent = Cm(0.74)
    para_after.paragraph_format.space_before = Pt(0)
    para_after.paragraph_format.space_after = Pt(0)
    para_after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    para_after.paragraph_format.line_spacing = Pt(12)

    doc.save(doc_path)



def has_images(folder_path: str) -> bool:
    return any(
        f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
        for f in os.listdir(folder_path)
    )



# 识别图片
def process_folder(folder_path: str,
                   log_callback=print,
                   use_typo_fix: bool = False,
                   use_editor: bool = False):
    folder_name = os.path.basename(folder_path)
    all_paragraphs = []

    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
            img_path = os.path.join(folder_path, filename)
            try:
                log_callback(f"正在识别：{img_path}")
                paragraphs = ocr_and_extract_text(img_path)
                all_paragraphs.extend(paragraphs)
            except Exception as e:
                log_callback(f"识别失败：{img_path}，原因：{e}")

    if all_paragraphs:
        doc_path = os.path.join(folder_path, f"{folder_name}.docx")
        log_callback(f"正在生成 Word：{doc_path}")

        # ① 先生成 Word（修改前 / 修改后 框架）
        create_word(doc_path, all_paragraphs, folder_name)

        # ② 再根据开关决定是否调用 LLM 纠错
        if use_typo_fix:
            log_callback("🤖 正在调用 LLM 进行错别字纠正...")
            try:
                fix_docx_with_llm(doc_path, task_name="typo_fix")
                log_callback("✅ LLM 纠错完成")
            except Exception as e:
                log_callback(f"⚠️ LLM 纠错失败：{e}")

        # ③ 如果启用了第二步编辑，再把纠正后的正文送到 editor 任务处理
        if use_editor:
            log_callback("🤖 正在调用 LLM 进行第二步作文改写...")
            try:
                edit_docx_with_llm(doc_path, task_name="editor")
                log_callback("✅ 第二步改写完成")
            except Exception as e:
                log_callback(f"⚠️ 第二步改写失败：{e}")

    else:
        log_callback(f"{folder_path} 中没有可识别的图片")

# 寻找‘修改前’与‘修改后’中 的文章
def extract_before_text(doc: Document):
    """Extract paragraphs between the marker "修改前：" and "修改后：".

    Returns a list of paragraph texts (str). If markers are missing returns an
    empty list. This function does not modify the document.
    """
    collecting = False
    paragraphs = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text == "修改前：":
            collecting = True
            continue

        if text == "修改后：":
            break

        if collecting and text:
            paragraphs.append(text)

    return paragraphs

# 清空「修改前」正文并写入新内容
def clear_before_text(doc: Document):
    """Remove all paragraphs between "修改前：" and "修改后：" (exclusive).

    Safe no-op if markers not found. Uses direct XML element removal which is
    supported by python-docx; callers should save the document after calling.
    """
    start = end = None

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "修改前：":
            start = i
        elif p.text.strip() == "修改后：" and start is not None:
            end = i
            break

    if start is None or end is None or end <= start:
        return

    # 删除 start+1 到 end-1
    for i in range(end - 1, start, -1):
        p = doc.paragraphs[i]
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            # 忽略单个段落删除错误，继续处理其它段落
            _LOGGER.exception("删除段落失败")





def insert_before_text(doc: Document, new_paragraphs):
    """
    在“修改前：”下面插入 AI 修正后的正文，并在“修改后：”之前插入分页符。
    如果没找到“修改后：”，则在刚插入的段落之后插入分页符（兜底）。
    """

    # 1) 找到“修改前：”所在位置（insert_index = 下一个段落的索引）
    insert_index = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "修改前：":
            insert_index = i + 1
            break
    if insert_index is None:
        _LOGGER.warning("未找到「修改前：」，跳过写入")
        return

    # 2) 记录已插入段落数，方便兜底时计算分页位置
    inserted_count = 0

    # 3) 插入 AI 返回的段落（保持顺序）
    for para in new_paragraphs:
        if not para or not para.strip():
            continue
        # 防止 AI 回传意外的标题
        if para.strip() in ("修改前：", "修改后："):
            continue

        new_p = doc.add_paragraph(para)

        # 设置格式：首行缩进、行距等
        fmt = new_p.paragraph_format
        fmt.first_line_indent = Cm(0.74)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        fmt.line_spacing = Pt(12)

        # 把刚创建在文档末尾的段落移动到指定位置
        try:
            doc._body._body.insert(insert_index, new_p._element)
        except Exception:
            _LOGGER.exception("插入段落失败")
            continue
        insert_index += 1
        inserted_count += 1

    # 4) 尝试在“修改后：”之前插入分页符（优先）
    page_break_index = None
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == "修改后：":
            page_break_index = idx
            break

    # 5) 如果没找到“修改后：”，把分页符放在插入段落之后（兜底）
    if page_break_index is None:
        page_break_index = insert_index  # 刚插入段落后的索引

    # 6) 创建一个新的段落并在其 run 上添加分页符（page break）
    #    先 add_paragraph（位于文档末尾），然后把它移动到目标位置
    pb_para = doc.add_paragraph()
    run = pb_para.add_run()
    run.add_break(WD_BREAK.PAGE)

    # 插入分页段落到目标位置（这样分页就在 page_break_index 处）
    try:
        doc._body._body.insert(page_break_index, pb_para._element)
    except Exception:
        _LOGGER.exception("插入分页符失败")
        # 兜底：追加到文档末尾
        pass


def insert_after_text(doc: Document, new_paragraphs):
    """
    在“修改后：”下面插入段落（不移动现有段落），如果未找到则追加到文档末尾。
    """
    insert_index = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "修改后：":
            insert_index = i + 1
            break
    if insert_index is None:
        insert_index = len(doc.paragraphs)

    for para in new_paragraphs:
        if not para or not para.strip():
            continue
        if para.strip() in ("修改前：", "修改后："):
            continue
        new_p = doc.add_paragraph(para)
        fmt = new_p.paragraph_format
        fmt.first_line_indent = Cm(0.74)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        fmt.line_spacing = Pt(12)
        try:
            doc._body._body.insert(insert_index, new_p._element)
        except Exception:
            _LOGGER.exception("insert_after_text 插入段落失败")
            continue
        insert_index += 1



# ai 纠错功能
def fix_docx_with_llm(docx_path: str, task_name: str = "typo_fix"):
    """Use an OpenAI-compatible LLM to fix typos inside the '修改前' section of the docx."""
    task_cfg = (config.get("LLM", {}) or {}).get("TASKS", {}).get(task_name, {})
    if not bool(task_cfg.get("ENABLED", False)):
        _LOGGER.info("LLM 纠错未启用，跳过")
        return

    doc = Document(docx_path)
    before_paragraphs = extract_before_text(doc)
    if not before_paragraphs:
        _LOGGER.info("未找到修改前正文，跳过")
        return

    full_text = "\n".join(before_paragraphs)

    clients, model, prompt_template = resolve_task_clients(config, task_name)
    if not prompt_template:
        prompt_template = DEFAULT_TYPO_FIX_PROMPT

    last_err = None
    fixed_text = None
    for base_url, client in clients:
        try:
            _LOGGER.info(f"尝试 LLM 节点：{base_url}")
            fixed_text = llm_fix_typos(prompt_template, full_text, client, model=model)
            _LOGGER.info(f"LLM 节点成功：{base_url}")
            break
        except Exception as e:
            last_err = e
            _LOGGER.warning(f"LLM 节点失败：{base_url} -> {e}")

    if fixed_text is None:
        raise last_err or RuntimeError("所有 LLM 节点均不可用")

    fixed_paragraphs = [p.strip() for p in fixed_text.split("\n") if p.strip()]

    clear_before_text(doc)
    insert_before_text(doc, fixed_paragraphs)
    doc.save(docx_path)


def edit_docx_with_llm(docx_path: str, task_name: str = "editor"):
    """使用 OpenAI-compatible LLM 对纠错后的文章进行进一步修改，并插入到“修改后：”之后。"""
    task_cfg = (config.get("LLM", {}) or {}).get("TASKS", {}).get(task_name, {})
    if not bool(task_cfg.get("ENABLED", False)):
        _LOGGER.info("第二步编辑未启用，跳过")
        return

    clients, model, prompt_template = resolve_task_clients(config, task_name)

    doc = Document(docx_path)

    corrected_paragraphs = extract_before_text(doc)
    if not corrected_paragraphs:
        _LOGGER.info("未找到可供第二步处理的正文，跳过")
        return

    full_text = "\n".join(corrected_paragraphs)

    if not prompt_template:
        prompt_template = DEFAULT_TYPO_FIX_PROMPT

    if "{text}" in prompt_template:
        prompt = prompt_template.format(text=full_text)
    else:
        prompt = prompt_template + "\n\n" + full_text

    if DEBUG:
        print("\n=================【发送给 editor LLM 的内容】=================\n")
        print(prompt)
        print("\n=============================================================\n")

    last_err = None
    result_text = None
    for base_url, client in clients:
        try:
            _LOGGER.info(f"尝试 editor LLM 节点：{base_url}")
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一名严谨的中文写作编辑助手"},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                stream=False,
            )

            if DEBUG:
                print("\n=================【editor LLM 原始返回 response】=================\n")
                try:
                    print(response)
                except Exception:
                    print(repr(response))
                print("\n===============================================================\n")

            result_text = response.choices[0].message.content.strip()
            _LOGGER.info(f"editor LLM 节点成功：{base_url}")
            break
        except Exception as e:
            last_err = e
            _LOGGER.warning(f"editor LLM 节点失败：{base_url} -> {e}")

    if result_text is None:
        raise last_err or RuntimeError("所有 editor LLM 节点均不可用")

    if DEBUG:
        print("\n=================【editor LLM 修改后的正文】=================\n")
        print(result_text)
        print("\n============================================================\n")

    edited_paragraphs = [p.strip() for p in result_text.split("\n") if p.strip()]
    insert_after_text(doc, edited_paragraphs)
    doc.save(docx_path)






# 遍历文件夹识别图片
def process_all(root_dir, log_callback=print, use_typo_fix=False, use_editor=False):
    if has_images(root_dir):
        process_folder(root_dir, log_callback, use_typo_fix=use_typo_fix, use_editor=use_editor)
    else:
        for sub in os.listdir(root_dir):
            sub_path = os.path.join(root_dir, sub)
            if os.path.isdir(sub_path) and has_images(sub_path):
                process_folder(sub_path, log_callback, use_typo_fix=use_typo_fix, use_editor=use_editor)


if __name__ == '__main__':
    ROOT_DIR = input("请输入要处理的文件夹路径：").strip('" ')
    if not os.path.isdir(ROOT_DIR):
        print("无效路径！")
    else:
        # For interactive usage, follow config toggles.
        tasks = (config.get("LLM", {}) or {}).get("TASKS", {})
        process_all(
            ROOT_DIR,
            use_typo_fix=bool(tasks.get("typo_fix", {}).get("ENABLED", False)),
            use_editor=bool(tasks.get("editor", {}).get("ENABLED", False)),
        )
        print("全部处理完成！")
