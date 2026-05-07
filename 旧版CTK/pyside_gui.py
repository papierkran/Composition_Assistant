# coding=utf-8
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from PySide6.QtCore import QObject, QThread, Signal
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QFileDialog,
    QFormLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QListWidget,
    QListWidgetItem,
    QLineEdit,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QSplitter,
    QStackedWidget,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)
from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openai import OpenAI

from config_migrate import ensure_new_schema

# OCR_CONFIG_FILE="D:\\person_data\\ocer助手\\presson.json"
CONFIG_FILE = Path(os.environ.get("OCR_CONFIG_FILE", "config.json")).expanduser()
# CONFIG_FILE = OCR_CONFIG_FILE


def load_config(path: Path = CONFIG_FILE) -> Dict[str, Any]:
    if not path.exists():
        local = Path("config.json")
        if local.exists():
            path = local
        else:
            return ensure_new_schema({})
    try:
        return ensure_new_schema(json.loads(path.read_text(encoding="utf-8")))
    except Exception:
        return ensure_new_schema({})


def save_config(cfg: Dict[str, Any], path: Path = CONFIG_FILE) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


class Worker(QObject):
    log = Signal(str)
    done = Signal()
    failed = Signal(str)

    def __init__(self, cfg: Dict[str, Any]):
        super().__init__()
        self.cfg = cfg

    def run(self):
        try:
            # 延迟导入：避免 GUI 启动时触发 ocr_main 顶层校验并直接崩溃
            from ocr_main import process_all

            root_dir = (self.cfg.get("APP", {}) or {}).get("ROOT_DIR", "").strip()
            tasks = (self.cfg.get("LLM", {}) or {}).get("TASKS", {})
            use_typo_fix = bool((tasks.get("typo_fix") or {}).get("ENABLED", False))
            use_editor = bool((tasks.get("editor") or {}).get("ENABLED", False))
            process_all(
                root_dir,
                log_callback=self.log.emit,
                use_typo_fix=use_typo_fix,
                use_editor=use_editor,
            )
            self.done.emit()
        except Exception as e:
            self.failed.emit(str(e))


class WorkflowWorker(QObject):
    log = Signal(str)
    done = Signal()
    failed = Signal(str)

    def __init__(self, folder: str, api_key: str, base_url: str, prompt: str, tasks: list[dict[str, Any]]):
        super().__init__()
        self.folder = folder
        self.api_key = api_key
        self.base_url = base_url
        self.prompt = prompt
        self.tasks = tasks

    def run(self):
        try:
            self._run_workflow()
            self.done.emit()
        except Exception as e:
            self.failed.emit(str(e))

    def _log(self, msg: str):
        self.log.emit(msg)

    def _iter_files_limited(self, folder: str, max_depth: int = 4):
        folder = os.path.abspath(folder)
        for root, dirs, files in os.walk(folder, topdown=True):
            rel = os.path.relpath(root, folder)
            depth = 0 if rel == os.curdir else len(rel.split(os.sep))
            if depth >= max_depth - 1:
                dirs[:] = []
            yield root, files

    def _run_workflow(self):
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif")
        copied_files = []
        self._log("【准备】复制原始文件（docx + 图片）...")
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if (name_lower.endswith(".docx") or name_lower.endswith(image_exts)) and not name_check.startswith("~$") and not name_check.startswith("改 "):
                    original_path = os.path.join(root, file)
                    new_path = os.path.join(root, f"改 {file}")
                    try:
                        shutil.copy2(original_path, new_path)
                        copied_files.append(new_path)
                    except Exception as e:
                        self._log(f"  ✗ 复制失败 {file}: {e}")
        if not copied_files:
            raise RuntimeError("未找到可处理的 docx 或图片文件")
        self._log(f"【准备】复制完成，共 {len(copied_files)} 个文件")

        ordered_enabled_tasks = sorted((t for t in self.tasks if t.get("enabled", True)), key=lambda x: x.get("order", 0))
        for task in ordered_enabled_tasks:
            task_id = task["id"]
            if task_id == "6":
                self._log("【步骤 6】转换 DOC -> DOCX")
                self._convert_docs()
            elif task_id == "1":
                self._log("【步骤 1】清除空格")
                self._clear_spaces()
            elif task_id == "AI":
                self._log("【步骤 AI】发送给 AI 修正")
                self._process_ai()
            elif task_id == "2":
                self._log("【步骤 2】添加标签")
                self._add_labels()
            elif task_id == "3":
                self._log("【步骤 3】格式化")
                self._format_docs()
            elif task_id == "5":
                self._log("【步骤 5】修改作者")
                self._set_author()

    def _convert_docs(self):
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 "):
                    continue
                if name_lower.endswith(".doc") and not name_check.startswith("~$"):
                    doc_path = os.path.join(root, file)
                    cmd = ["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", root]
                    try:
                        subprocess.run(cmd, capture_output=True, timeout=30)
                    except Exception as e:
                        self._log(f"  ✗ DOC 转换失败 {file}: {e}")

    def _clear_spaces(self):
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 "):
                    continue
                if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                    doc_path = os.path.join(root, file)
                    try:
                        doc = Document(doc_path)
                        for para in doc.paragraphs:
                            for run in para.runs:
                                run.text = run.text.strip()
                        doc.save(doc_path)
                    except Exception as e:
                        self._log(f"  ✗ 清空格失败 {file}: {e}")

    def _process_ai(self):
        client = OpenAI(api_key=self.api_key, base_url=self.base_url)
        prompt_template = self.prompt or "{text}"
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 ") or not name_lower.endswith(".docx") or name_check.startswith("~$"):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    all_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    if not all_text.strip():
                        continue
                    full_prompt = prompt_template.format(text=all_text) if "{text}" in prompt_template else f"{prompt_template}\n\n{all_text}"
                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[
                            {"role": "system", "content": "你是一名严谨的中文校对助手"},
                            {"role": "user", "content": full_prompt},
                        ],
                        temperature=0.1,
                        stream=False,
                    )
                    ai_result = response.choices[0].message.content.strip()
                    last_para = doc.paragraphs[-1] if doc.paragraphs else None
                    if last_para:
                        if last_para.runs:
                            last_para.runs[-1].add_break(WD_BREAK.PAGE)
                        else:
                            last_para.add_run().add_break(WD_BREAK.PAGE)
                    para_modify_label = doc.add_paragraph("修改后：")
                    para_modify_label.paragraph_format.first_line_indent = Cm(0.74)
                    para_modify_label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    para_modify_label.paragraph_format.line_spacing = Pt(12)
                    for line in ai_result.split("\n"):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            fmt = p.paragraph_format
                            fmt.first_line_indent = Cm(0.74)
                            fmt.space_before = Pt(0)
                            fmt.space_after = Pt(0)
                            fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            fmt.line_spacing = Pt(12)
                    doc.save(doc_path)
                except Exception as e:
                    self._log(f"  ✗ AI 处理失败 {file}: {e}")

    def _add_labels(self):
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 ") or not name_lower.endswith(".docx") or name_check.startswith("~$"):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    if doc.paragraphs and doc.paragraphs[0].text.strip() != "修改前：":
                        doc.paragraphs[0].insert_paragraph_before("修改前：")
                    doc.save(doc_path)
                except Exception as e:
                    self._log(f"  ✗ 添加标签失败 {file}: {e}")

    def _format_docs(self):
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 ") or not name_lower.endswith(".docx") or name_check.startswith("~$"):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    style = doc.styles["Normal"]
                    style.font.name = "宋体"
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    style.font.size = Pt(12)
                    for para in doc.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0.74)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                except Exception as e:
                    self._log(f"  ✗ 格式化失败 {file}: {e}")

    def _set_author(self):
        for root, files in self._iter_files_limited(self.folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if not name_check.startswith("改 ") or not name_lower.endswith(".docx") or name_check.startswith("~$"):
                    continue
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    doc.core_properties.author = "思睿教育_美丽可爱的尹老师"
                    doc.save(doc_path)
                except Exception as e:
                    self._log(f"  ✗ 修改作者失败 {file}: {e}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Composition OCR Assistant - Qt")
        self.resize(1180, 760)
        self.config = load_config()
        self.thread: QThread | None = None
        self.worker: Worker | None = None
        self.workflow_thread: QThread | None = None
        self.workflow_worker: WorkflowWorker | None = None

        root = QWidget()
        root_layout = QHBoxLayout(root)
        root_layout.setContentsMargins(8, 8, 8, 8)

        self.nav = QListWidget()
        self.nav.setFixedWidth(180)
        for title in ["图片转作文", "docx作文处理", "配置编辑"]:
            QListWidgetItem(title, self.nav)
        self.nav.setCurrentRow(0)

        self.stack = QStackedWidget()
        self.stack.addWidget(self._build_ocr_page())
        self.stack.addWidget(self._build_workflow_page())
        self.stack.addWidget(self._build_config_page())
        self.nav.currentRowChanged.connect(self.stack.setCurrentIndex)

        root_layout.addWidget(self.nav)
        root_layout.addWidget(self.stack, 1)
        self.setCentralWidget(root)
        self._refresh_ui_from_config()

    def _build_ocr_page(self) -> QWidget:
        page = QWidget()
        layout = QVBoxLayout(page)

        splitter = QSplitter()
        left = QWidget()
        right = QWidget()
        splitter.addWidget(left)
        splitter.addWidget(right)
        splitter.setSizes([620, 560])
        layout.addWidget(splitter)

        left_layout = QVBoxLayout(left)
        right_layout = QVBoxLayout(right)

        run_box = QGroupBox("处理设置")
        run_form = QFormLayout(run_box)

        self.root_dir_edit = QLineEdit()
        browse_btn = QPushButton("浏览")
        browse_btn.clicked.connect(self._browse_dir)
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        row_layout.addWidget(self.root_dir_edit)
        row_layout.addWidget(browse_btn)
        run_form.addRow("作文文件夹", row)

        self.typo_enabled = QCheckBox("启用 AI 错别字修正")
        self.editor_enabled = QCheckBox("启用第二步改写")
        run_form.addRow("", self.typo_enabled)
        run_form.addRow("", self.editor_enabled)

        self.start_btn = QPushButton("开始处理")
        self.start_btn.clicked.connect(self._start)
        run_form.addRow("", self.start_btn)
        left_layout.addWidget(run_box)

        provider_box = QGroupBox("基础 OCR / Provider 配置")
        provider_form = QFormLayout(provider_box)
        self.ocr_url = QLineEdit()
        self.ocr_appid = QLineEdit()
        self.ocr_key = QLineEdit()
        self.ocr_key.setEchoMode(QLineEdit.Password)
        provider_form.addRow("OCR URL", self.ocr_url)
        provider_form.addRow("OCR APPID", self.ocr_appid)
        provider_form.addRow("OCR API_KEY", self.ocr_key)

        self.typo_provider = QComboBox()
        self.typo_provider.addItems(["deepseek", "openai", "custom"])
        self.typo_prompt = QTextEdit()
        self.typo_prompt.setMinimumHeight(120)
        provider_form.addRow("错别字任务 Provider", self.typo_provider)
        provider_form.addRow("错别字 Prompt", self.typo_prompt)
        left_layout.addWidget(provider_box)

        log_box = QGroupBox("运行日志")
        log_layout = QVBoxLayout(log_box)
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        log_layout.addWidget(self.log_text)
        right_layout.addWidget(log_box)

        return page

    def _build_workflow_page(self) -> QWidget:
        page = QWidget()
        layout = QVBoxLayout(page)
        splitter = QSplitter()
        left = QWidget()
        right = QWidget()
        splitter.addWidget(left)
        splitter.addWidget(right)
        splitter.setSizes([560, 620])
        layout.addWidget(splitter)

        left_layout = QVBoxLayout(left)
        right_layout = QVBoxLayout(right)

        box = QGroupBox("docx 处理流程")
        form = QFormLayout(box)
        self.workflow_folder = QLineEdit()
        wf_browse = QPushButton("浏览")
        wf_browse.clicked.connect(self._browse_workflow_dir)
        folder_row = QWidget()
        folder_row_layout = QHBoxLayout(folder_row)
        folder_row_layout.setContentsMargins(0, 0, 0, 0)
        folder_row_layout.addWidget(self.workflow_folder)
        folder_row_layout.addWidget(wf_browse)
        form.addRow("处理文件夹", folder_row)

        self.workflow_api_key = QLineEdit()
        self.workflow_api_key.setEchoMode(QLineEdit.Password)
        self.workflow_base_url = QLineEdit()
        self.workflow_prompt = QTextEdit()
        self.workflow_prompt.setMinimumHeight(100)
        form.addRow("AI API Key", self.workflow_api_key)
        form.addRow("API Base URL", self.workflow_base_url)
        form.addRow("AI Prompt", self.workflow_prompt)
        left_layout.addWidget(box)

        task_box = QGroupBox("任务顺序（可勾选 + 上下移动）")
        task_layout = QVBoxLayout(task_box)
        self.workflow_tasks = [
            {"id": "6", "name": "6. 转换 DOC -> DOCX", "enabled": True, "order": 0},
            {"id": "1", "name": "1. 清除空格", "enabled": True, "order": 1},
            {"id": "AI", "name": "AI. AI 改作文", "enabled": True, "order": 2},
            {"id": "2", "name": "2. 添加标签", "enabled": True, "order": 3},
            {"id": "3", "name": "3. 格式化", "enabled": True, "order": 4},
            {"id": "5", "name": "5. 修改作者", "enabled": True, "order": 5},
        ]
        self.workflow_task_list = QListWidget()
        self._refresh_workflow_task_list()
        self.workflow_task_list.itemDoubleClicked.connect(self._toggle_task_enabled)
        task_layout.addWidget(self.workflow_task_list)
        controls = QHBoxLayout()
        up_btn = QPushButton("上移")
        down_btn = QPushButton("下移")
        up_btn.clicked.connect(lambda: self._move_task(-1))
        down_btn.clicked.connect(lambda: self._move_task(1))
        controls.addWidget(up_btn)
        controls.addWidget(down_btn)
        task_layout.addLayout(controls)
        left_layout.addWidget(task_box)

        self.workflow_start_btn = QPushButton("开始 docx 流程")
        self.workflow_start_btn.clicked.connect(self._start_workflow)
        left_layout.addWidget(self.workflow_start_btn)

        workflow_log_box = QGroupBox("流程日志")
        workflow_log_layout = QVBoxLayout(workflow_log_box)
        self.workflow_log_text = QTextEdit()
        self.workflow_log_text.setReadOnly(True)
        workflow_log_layout.addWidget(self.workflow_log_text)
        right_layout.addWidget(workflow_log_box)
        return page

    def _build_config_page(self) -> QWidget:
        page = QWidget()
        layout = QVBoxLayout(page)

        group = QGroupBox("LLM Providers")
        form = QFormLayout(group)
        self.deepseek_base = QLineEdit()
        self.deepseek_model = QLineEdit()
        self.deepseek_key = QLineEdit()
        self.deepseek_key.setEchoMode(QLineEdit.Password)
        self.openai_base = QLineEdit()
        self.openai_model = QLineEdit()
        self.openai_key = QLineEdit()
        self.openai_key.setEchoMode(QLineEdit.Password)
        self.custom_base = QLineEdit()
        self.custom_model = QLineEdit()
        self.custom_key = QLineEdit()
        self.custom_key.setEchoMode(QLineEdit.Password)

        form.addRow("deepseek BASE_URL", self.deepseek_base)
        form.addRow("deepseek MODEL", self.deepseek_model)
        form.addRow("deepseek API_KEY", self.deepseek_key)
        form.addRow("openai BASE_URL", self.openai_base)
        form.addRow("openai MODEL", self.openai_model)
        form.addRow("openai API_KEY", self.openai_key)
        form.addRow("custom BASE_URL", self.custom_base)
        form.addRow("custom MODEL", self.custom_model)
        form.addRow("custom API_KEY", self.custom_key)

        layout.addWidget(group)
        save_btn = QPushButton("保存配置")
        save_btn.clicked.connect(self._save_from_ui)
        layout.addWidget(save_btn)
        layout.addStretch(1)
        return page

    def _refresh_ui_from_config(self):
        cfg = ensure_new_schema(self.config)
        self.config = cfg
        ocr = ((cfg.get("OCR") or {}).get("XFYUN") or {})
        app = cfg.get("APP") or {}
        llm = cfg.get("LLM") or {}
        providers = llm.get("PROVIDERS") or {}
        tasks = llm.get("TASKS") or {}

        self.root_dir_edit.setText(app.get("ROOT_DIR", ""))
        self.typo_enabled.setChecked(bool((tasks.get("typo_fix") or {}).get("ENABLED", False)))
        self.editor_enabled.setChecked(bool((tasks.get("editor") or {}).get("ENABLED", False)))
        self.ocr_url.setText(ocr.get("URL", ""))
        self.ocr_appid.setText(ocr.get("APPID", ""))
        self.ocr_key.setText(ocr.get("API_KEY", ""))
        self.typo_provider.setCurrentText((tasks.get("typo_fix") or {}).get("PROVIDER", "deepseek"))
        self.typo_prompt.setPlainText((tasks.get("typo_fix") or {}).get("PROMPT", "{text}"))

        deepseek = providers.get("deepseek") or {}
        openai = providers.get("openai") or {}
        custom = providers.get("custom") or {}
        self.deepseek_base.setText(deepseek.get("BASE_URL", ""))
        self.deepseek_model.setText(deepseek.get("MODEL", "deepseek-chat"))
        self.deepseek_key.setText(deepseek.get("API_KEY", ""))
        self.openai_base.setText(openai.get("BASE_URL", "https://api.openai.com/v1"))
        self.openai_model.setText(openai.get("MODEL", "gpt-4o-mini"))
        self.openai_key.setText(openai.get("API_KEY", ""))
        self.custom_base.setText(custom.get("BASE_URL", ""))
        self.custom_model.setText(custom.get("MODEL", ""))
        self.custom_key.setText(custom.get("API_KEY", ""))
        self.workflow_folder.setText(app.get("ROOT_DIR", ""))
        self.workflow_api_key.setText(custom.get("API_KEY", ""))
        self.workflow_base_url.setText(custom.get("BASE_URL", ""))
        self.workflow_prompt.setPlainText((tasks.get("editor") or {}).get("PROMPT", "{text}"))

    def _collect_ui_to_config(self) -> Dict[str, Any]:
        cfg = ensure_new_schema(self.config)
        cfg.setdefault("APP", {})
        cfg["APP"]["ROOT_DIR"] = self.root_dir_edit.text().strip()

        cfg.setdefault("OCR", {})
        cfg["OCR"]["PROVIDER"] = "xfyun_handwriting"
        cfg["OCR"].setdefault("XFYUN", {})
        cfg["OCR"]["XFYUN"]["URL"] = self.ocr_url.text().strip()
        cfg["OCR"]["XFYUN"]["APPID"] = self.ocr_appid.text().strip()
        cfg["OCR"]["XFYUN"]["API_KEY"] = self.ocr_key.text().strip()
        cfg["OCR"]["XFYUN"]["LANGUAGE"] = cfg["OCR"]["XFYUN"].get("LANGUAGE", "cn|en")
        cfg["OCR"]["XFYUN"]["LOCATION"] = cfg["OCR"]["XFYUN"].get("LOCATION", "false")

        cfg.setdefault("LLM", {})
        cfg["LLM"].setdefault("PROVIDERS", {})
        cfg["LLM"].setdefault("TASKS", {})

        cfg["LLM"]["PROVIDERS"]["deepseek"] = {
            "BASE_URL": self.deepseek_base.text().strip(),
            "MODEL": self.deepseek_model.text().strip() or "deepseek-chat",
            "API_KEY": self.deepseek_key.text().strip(),
        }
        cfg["LLM"]["PROVIDERS"]["openai"] = {
            "BASE_URL": self.openai_base.text().strip(),
            "MODEL": self.openai_model.text().strip() or "gpt-4o-mini",
            "API_KEY": self.openai_key.text().strip(),
        }
        cfg["LLM"]["PROVIDERS"]["custom"] = {
            "BASE_URL": self.custom_base.text().strip(),
            "MODEL": self.custom_model.text().strip(),
            "API_KEY": self.custom_key.text().strip(),
        }
        cfg["LLM"]["TASKS"]["typo_fix"] = {
            "ENABLED": bool(self.typo_enabled.isChecked()),
            "PROVIDER": self.typo_provider.currentText().strip() or "deepseek",
            "PROMPT": self.typo_prompt.toPlainText().strip() or "{text}",
        }
        cfg["LLM"]["TASKS"].setdefault("editor", {})
        cfg["LLM"]["TASKS"]["editor"]["ENABLED"] = bool(self.editor_enabled.isChecked())
        cfg["LLM"]["TASKS"]["editor"].setdefault("PROVIDER", "deepseek")
        cfg["LLM"]["TASKS"]["editor"].setdefault("PROMPT", "{text}")
        return cfg

    def _save_from_ui(self):
        cfg = self._collect_ui_to_config()
        if not cfg["OCR"]["XFYUN"]["URL"] or not cfg["OCR"]["XFYUN"]["APPID"] or not cfg["OCR"]["XFYUN"]["API_KEY"]:
            QMessageBox.warning(self, "配置不完整", "OCR 的 URL / APPID / API_KEY 必填。")
            return
        save_config(cfg)
        self.config = cfg
        QMessageBox.information(self, "保存成功", f"配置已写入：{CONFIG_FILE}")

    def _browse_dir(self):
        path = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if path:
            self.root_dir_edit.setText(path)

    def _browse_workflow_dir(self):
        path = QFileDialog.getExistingDirectory(self, "选择 docx 处理目录")
        if path:
            self.workflow_folder.setText(path)

    def _append_log(self, msg: str):
        now = datetime.now().strftime("%H:%M:%S")
        self.log_text.append(f"[{now}] {msg}")

    def _append_workflow_log(self, msg: str):
        now = datetime.now().strftime("%H:%M:%S")
        self.workflow_log_text.append(f"[{now}] {msg}")

    def _refresh_workflow_task_list(self):
        self.workflow_task_list.clear()
        self.workflow_tasks.sort(key=lambda x: x["order"])
        for task in self.workflow_tasks:
            item = QListWidgetItem(f"{'☑' if task['enabled'] else '☐'} {task['name']}")
            item.setData(32, task["id"])
            self.workflow_task_list.addItem(item)

    def _move_task(self, delta: int):
        row = self.workflow_task_list.currentRow()
        if row < 0:
            return
        target = row + delta
        if target < 0 or target >= len(self.workflow_tasks):
            return
        self.workflow_tasks[row], self.workflow_tasks[target] = self.workflow_tasks[target], self.workflow_tasks[row]
        for idx, task in enumerate(self.workflow_tasks):
            task["order"] = idx
        self._refresh_workflow_task_list()
        self.workflow_task_list.setCurrentRow(target)

    def _toggle_task_enabled(self, item: QListWidgetItem):
        task_id = item.data(32)
        for task in self.workflow_tasks:
            if task["id"] == task_id:
                task["enabled"] = not bool(task.get("enabled", True))
                break
        self._refresh_workflow_task_list()

    def _start_workflow(self):
        folder = self.workflow_folder.text().strip()
        api_key = self.workflow_api_key.text().strip()
        base_url = self.workflow_base_url.text().strip()
        prompt = self.workflow_prompt.toPlainText().strip() or "{text}"
        if not folder or not Path(folder).exists():
            QMessageBox.warning(self, "路径错误", "请先选择有效的 docx 处理目录。")
            return
        if not api_key:
            QMessageBox.warning(self, "配置缺失", "docx 流程需要 AI API Key。")
            return

        self.workflow_start_btn.setEnabled(False)
        self._append_workflow_log("🚀 开始执行 docx 流程")
        self.workflow_thread = QThread(self)
        tasks_snapshot = [dict(t) for t in self.workflow_tasks]
        self.workflow_worker = WorkflowWorker(folder, api_key, base_url, prompt, tasks_snapshot)
        self.workflow_worker.moveToThread(self.workflow_thread)
        self.workflow_thread.started.connect(self.workflow_worker.run)
        self.workflow_worker.log.connect(self._append_workflow_log)
        self.workflow_worker.done.connect(self._on_workflow_done)
        self.workflow_worker.failed.connect(self._on_workflow_failed)
        self.workflow_worker.done.connect(self.workflow_thread.quit)
        self.workflow_worker.failed.connect(self.workflow_thread.quit)
        self.workflow_thread.finished.connect(self.workflow_thread.deleteLater)
        self.workflow_thread.start()

    def _start(self):
        cfg = self._collect_ui_to_config()
        root_dir = (cfg.get("APP", {}) or {}).get("ROOT_DIR", "")
        if not root_dir or not Path(root_dir).exists():
            QMessageBox.warning(self, "路径错误", "请先选择有效的作文文件夹路径。")
            return
        try:
            save_config(cfg)
        except Exception as e:
            QMessageBox.critical(self, "保存失败", str(e))
            return
        self.config = cfg

        self.start_btn.setEnabled(False)
        self._append_log("🚀 开始处理")

        self.thread = QThread(self)
        self.worker = Worker(cfg)
        self.worker.moveToThread(self.thread)
        self.thread.started.connect(self.worker.run)
        self.worker.log.connect(self._append_log)
        self.worker.done.connect(self._on_done)
        self.worker.failed.connect(self._on_failed)
        self.worker.done.connect(self.thread.quit)
        self.worker.failed.connect(self.thread.quit)
        self.thread.finished.connect(self.thread.deleteLater)
        self.thread.start()

    def _on_done(self):
        self._append_log("✅ 全部处理完成")
        self.start_btn.setEnabled(True)

    def _on_failed(self, err: str):
        self._append_log(f"❌ 处理失败：{err}")
        self.start_btn.setEnabled(True)

    def _on_workflow_done(self):
        self._append_workflow_log("✅ docx 流程执行完成")
        self.workflow_start_btn.setEnabled(True)

    def _on_workflow_failed(self, err: str):
        self._append_workflow_log(f"❌ docx 流程失败：{err}")
        self.workflow_start_btn.setEnabled(True)


def main():
    app = QApplication(sys.argv)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
