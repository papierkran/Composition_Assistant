import base64
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import json
import threading
from datetime import datetime
import sys
import os
import subprocess
import re
from pathlib import Path
from copy import deepcopy
import customtkinter as ctk
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from openai import OpenAI
from PIL import Image, ImageTk
import ctypes
import io


# 默认使用项目目录下的 config.json；如需覆盖，可通过环境变量 OCR_CONFIG_FILE 指定。
CONFIG_FILE = Path(os.environ.get("OCR_CONFIG_FILE", "config.json")).expanduser()

# ================= 默认配置（新结构） =================
DEFAULT_CONFIG = {
    "OCR": {
        "PROVIDER": "xfyun_handwriting",
        "XFYUN": {
            "URL": "http://webapi.xfyun.cn/v1/service/v1/ocr/handwriting",
            "APPID": "",
            "API_KEY": "",
            "LANGUAGE": "cn|en",
            "LOCATION": "false",
        },
    },
    "LLM": {
        "PROVIDERS": {
            "deepseek": {
                "API_KEY": "",
                "MODEL": "deepseek-chat",
                "BASE_URL": "https://api.deepseek.com/v1",
            },
            "openai": {
                "API_KEY": "",
                "MODEL": "gpt-4o-mini",
                "BASE_URL": "https://api.openai.com/v1",
            },
            "custom": {
                "API_KEY": "",
                "MODEL": "",
                "BASE_URL": "",
            },
        },
        "TASKS": {
            "typo_fix": {
                "ENABLED": False,
                "PROVIDER": "deepseek",
                "PROMPT": "{text}",
            },
            "editor": {
                "ENABLED": False,
                "PROVIDER": "deepseek",
                "PROMPT": "{text}",
                "COUNT_MIN": None,
                "COUNT_MAX": None,
            },
        },
    },
    "APP": {
        "ROOT_DIR": "",
        "DEBUG": False
    }
}

# ================= 配置文件 =================
from config_migrate import ensure_new_schema


def load_config(path: Path = None):
    """
    从 JSON 文件加载配置。
    尝试使用 `path`（或配置的 CONFIG_FILE）。
    如果缺失，则回退到 './config.json'。在任何读取错误的情况下，返回 DEFAULT_CONFIG 的浅拷贝以允许界面启动。
    """
    cfg_path = Path(path or CONFIG_FILE)
    if not cfg_path.exists():
        local = Path("config.json")
        if local.exists():
            cfg_path = local
        else:
            return deepcopy(DEFAULT_CONFIG)

    try:
        with cfg_path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return deepcopy(DEFAULT_CONFIG)

def save_config(config, path: Path = None):
    """Save configuration to file. Ensures parent directory exists."""
    cfg_path = Path(path or CONFIG_FILE)
    cfg_path.parent.mkdir(parents=True, exist_ok=True)
    with cfg_path.open("w", encoding="utf-8") as f:
        json.dump(config, f, indent=2, ensure_ascii=False)

# ================= 日志 =================
def append_log(message: str):
    timestamp = datetime.now().strftime("%H:%M:%S")
    log_text.configure(state="normal")
    log_text.insert(tk.END, f"[{timestamp}] {message}\n")
    log_text.see(tk.END)
    log_text.configure(state="disabled")

# ================= 配置编辑窗口（表单） =================
from config_editor_ui import open_config_editor_form


def open_config_editor():
    def _on_saved(new_cfg):
        # 更新内存 config，并尽量刷新主页面输入框
        global config
        config = ensure_new_schema(new_cfg)
        try:
            url_entry.delete(0, tk.END)
            url_entry.insert(0, config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
            appid_entry.delete(0, tk.END)
            appid_entry.insert(0, config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
            path_entry.delete(0, tk.END)
            path_entry.insert(0, config.get("APP", {}).get("ROOT_DIR", ""))

            # 开关/提示词
            use_deepseek_var.set(bool((config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
            prompt_text.delete("1.0", tk.END)
            prompt_text.insert("1.0", (config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT", "{text}"))

            use_editor_var.set(bool((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
            editor_prompt_text.delete("1.0", tk.END)
            editor_prompt_text.insert("1.0", (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"))

            # 刷新 provider 下拉，避免删除后主界面仍显示旧 provider
            names = _provider_name_list()
            typo_provider = _normalize_provider_name(
                (config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROVIDER", "deepseek")
            ) or "deepseek"
            edit_provider = _normalize_provider_name(
                (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")
            ) or "deepseek"
            if typo_provider not in names:
                typo_provider = names[0]
            if edit_provider not in names:
                edit_provider = names[0]

            deepseek_provider_menu.configure(values=names)
            deepseek_provider_menu.set(typo_provider)
            _on_deepseek_provider_change(typo_provider)

            editor_provider_menu.configure(values=names)
            editor_provider_menu.set(edit_provider)
            _on_editor_provider_change(edit_provider)

            if "ai_provider_menu" in globals():
                ai_provider_menu.configure(values=names)
                ai_provider = edit_provider if edit_provider in names else names[0]
                ai_provider_menu.set(ai_provider)
                _on_ai_provider_change(ai_provider)
        except Exception:
            pass

    open_config_editor_form(
        root=root,
        config=config,
        config_file=CONFIG_FILE,
        hidden_api_keys=hidden_api_keys,
        make_mask_widget=make_mask_widget,
        reveal_entry=reveal_entry,
        on_saved=_on_saved,
    )


# ================= 主逻辑 =================
def start_processing():
    def task():
        # ---------- 从 UI 读取（新 schema） ----------
        config.setdefault("OCR", {})
        config["OCR"].setdefault("XFYUN", {})
        config["OCR"]["PROVIDER"] = "xfyun_handwriting"
        config["OCR"]["XFYUN"]["URL"] = url_entry.get().strip()
        config["OCR"]["XFYUN"]["APPID"] = appid_entry.get().strip()
        config["OCR"]["XFYUN"]["API_KEY"] = get_api_key_value('ocr') or ""
        config["OCR"]["XFYUN"]["LANGUAGE"] = config["OCR"]["XFYUN"].get("LANGUAGE", "cn|en")
        config["OCR"]["XFYUN"]["LOCATION"] = config["OCR"]["XFYUN"].get("LOCATION", "false")

        config.setdefault("APP", {})
        config["APP"]["ROOT_DIR"] = path_entry.get().strip()

        config.setdefault("LLM", {})
        config["LLM"].setdefault("PROVIDERS", {})
        config["LLM"].setdefault("TASKS", {})

        # typo_fix task/provider
        typo_provider = _ensure_provider_exists(deepseek_provider_menu.get() or "deepseek")

        config["LLM"]["PROVIDERS"].setdefault(typo_provider, {})
        config["LLM"]["PROVIDERS"][typo_provider]["API_KEY"] = get_api_key_value('deepseek') or ""
        config["LLM"]["PROVIDERS"][typo_provider]["BASE_URL"] = deepseek_base_entry.get().strip()
        config["LLM"]["PROVIDERS"][typo_provider].setdefault("MODEL", "deepseek-chat" if typo_provider == "deepseek" else "gpt-4o-mini")

        config["LLM"]["TASKS"].setdefault("typo_fix", {})
        config["LLM"]["TASKS"]["typo_fix"]["ENABLED"] = use_deepseek_var.get()
        config["LLM"]["TASKS"]["typo_fix"]["PROVIDER"] = typo_provider
        config["LLM"]["TASKS"]["typo_fix"]["PROMPT"] = prompt_text.get("1.0", tk.END).strip()

        # editor task/provider
        editor_provider = _ensure_provider_exists(editor_provider_menu.get() or "deepseek")

        config["LLM"]["PROVIDERS"].setdefault(editor_provider, {})
        config["LLM"]["PROVIDERS"][editor_provider]["API_KEY"] = get_api_key_value('editor') or ""
        config["LLM"]["PROVIDERS"][editor_provider]["BASE_URL"] = editor_base_entry.get().strip()
        config["LLM"]["PROVIDERS"][editor_provider].setdefault("MODEL", "deepseek-chat" if editor_provider == "deepseek" else "gpt-4o-mini")

        count_min_text = editor_count_min_entry.get().strip()
        count_max_text = editor_count_max_entry.get().strip()
        count_min = None
        count_max = None
        if count_min_text:
            try:
                count_min = int(count_min_text)
            except ValueError:
                append_log("❌ 目标字数范围最小值必须是整数")
                return
        if count_max_text:
            try:
                count_max = int(count_max_text)
            except ValueError:
                append_log("❌ 目标字数范围最大值必须是整数")
                return
        if count_min is not None and count_max is not None and count_min > count_max:
            append_log("❌ 目标字数范围最小值不能大于最大值")
            return

        config["LLM"]["TASKS"].setdefault("editor", {})
        config["LLM"]["TASKS"]["editor"]["ENABLED"] = use_editor_var.get()
        config["LLM"]["TASKS"]["editor"]["PROVIDER"] = editor_provider
        config["LLM"]["TASKS"]["editor"]["PROMPT"] = editor_prompt_text.get("1.0", tk.END).strip()
        config["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        config["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max

        # ---------- 百度图片矫正配置 ----------
        config["OCR"].setdefault("BAIDU_CORRECTION", {})
        config["OCR"]["BAIDU_CORRECTION"]["ENABLED"] = use_baidu_correction_var.get()
        config["OCR"]["BAIDU_CORRECTION"]["API_KEY"] = get_api_key_value('baidu_api_key') or ""
        config["OCR"]["BAIDU_CORRECTION"]["SECRET_KEY"] = get_api_key_value('baidu_secret_key') or ""

        # ---------- 校验 ----------
        if not all([
            config.get("OCR", {}).get("XFYUN", {}).get("URL"),
            config.get("OCR", {}).get("XFYUN", {}).get("APPID"),
            config.get("OCR", {}).get("XFYUN", {}).get("API_KEY"),
            config.get("APP", {}).get("ROOT_DIR"),
        ]):
            append_log("❌ 请填写完整的 OCR 配置和文件夹路径")
            return

        # 校验：启用图片矫正时必须有百度 API Key 和 Secret Key
        if bool(config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED", False)):
            if not config["OCR"]["BAIDU_CORRECTION"].get("API_KEY") or not config["OCR"]["BAIDU_CORRECTION"].get("SECRET_KEY"):
                append_log("❌ 已启用图片矫正，但未填写百度 API Key 或 Secret Key")
                return

        # 校验：启用 typo_fix 时必须有对应 provider 的 API Key
        tasks = (config.get("LLM", {}) or {}).get("TASKS", {})
        providers = (config.get("LLM", {}) or {}).get("PROVIDERS", {})
        if bool(tasks.get("typo_fix", {}).get("ENABLED", False)):
            p_name = tasks.get("typo_fix", {}).get("PROVIDER")
            if not (providers.get(p_name, {}) or {}).get("API_KEY"):
                append_log("❌ 已启用 AI 错别字修正，但未填写 API Key")
                return

        if not os.path.isdir(config["APP"]["ROOT_DIR"]):
            append_log("❌ 文件夹路径无效")
            return

        save_config(config)

        # ---------- 日志 ----------
        if bool(tasks.get("typo_fix", {}).get("ENABLED", False)):
            append_log("🧠 AI 错别字修正：已启用")
        else:
            append_log("🧠 AI 错别字修正：未启用")

        if bool(config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED", False)):
            append_log("📷 百度图片矫正：已启用")
        else:
            append_log("📷 百度图片矫正：未启用")

        def task_status_callback(folder_path: str, status: str):
            root.after(0, lambda: update_task_status(folder_path, status))

        try:
            append_log("🚀 开始处理...")
            from ocr_main import process_all

            process_all(
                config["APP"]["ROOT_DIR"],
                log_callback=append_log,
                use_typo_fix=bool(tasks.get("typo_fix", {}).get("ENABLED", False)),
                use_editor=bool(tasks.get("editor", {}).get("ENABLED", False)),
                cfg=config,  # 传入当前配置以更新 ocr_main 的全局变量
                task_status_callback=task_status_callback,
            )

            append_log("✅ 全部处理完成")
        except Exception as e:
            append_log(f"❌ 处理失败：{e}")

    refresh_task_queue()
    threading.Thread(target=task, daemon=True).start()

# ================= 选择文件夹 =================
def browse_folder():
    folder = filedialog.askdirectory()
    if folder:
        path_entry.delete(0, tk.END)
        path_entry.insert(0, folder)
        refresh_task_queue()


def iter_files_limited(folder, max_depth=4):
    """遍历 folder，最多递归 max_depth 层（包含根目录为第0层）。
    返回 (root, files) 与 os.walk 类似，但在达到深度限制时不再向下递归。
    """
    folder = os.path.abspath(folder)
    for root, dirs, files in os.walk(folder, topdown=True):
        rel = os.path.relpath(root, folder)
        if rel == os.curdir:
            depth = 0
        else:
            depth = len(rel.split(os.sep))
        # 深度为 0..(max_depth-1) 可被处理；当达到阈值时阻止继续向下遍历
        if depth >= max_depth - 1:
            dirs[:] = []
        yield root, files

# ================= UI =================
config = ensure_new_schema(load_config())

# 存储已隐藏的 API Key
hidden_api_keys = {}
# 当前活跃的 entry 映射（name->entry 或 None）
entries_map = {}

# 预置 Provider 模板（可快速填充 BASE_URL）
API_PROVIDER_PRESETS = {
    "deepseek": "https://api.deepseek.com/v1",
    "openai": "https://api.openai.com/v1",
}


def _normalize_provider_name(name: str) -> str:
    return (name or "").strip().lower()


def _ensure_provider_exists(name: str):
    p_name = _normalize_provider_name(name)
    if not p_name:
        return ""
    config.setdefault("LLM", {}).setdefault("PROVIDERS", {}).setdefault(p_name, {})
    p_cfg = config["LLM"]["PROVIDERS"][p_name]
    p_cfg.setdefault("API_KEY", "")
    p_cfg.setdefault("BASE_URL", API_PROVIDER_PRESETS.get(p_name, ""))
    p_cfg.setdefault("MODEL", "deepseek-chat" if p_name == "deepseek" else "gpt-4o-mini")
    return p_name


def _provider_name_list():
    providers = (config.get("LLM", {}) or {}).get("PROVIDERS", {}) or {}
    names = {k for k in providers.keys() if isinstance(k, str) and k.strip()}
    names.update(API_PROVIDER_PRESETS.keys())
    return sorted(names)


def _open_new_provider_dialog(on_confirm):
    win = ctk.CTkToplevel(root)
    win.title("新增 AI Provider")
    win.geometry("550x240")
    try:
        win.transient(root)
        win.grab_set()
    except Exception:
        pass

    frm = ctk.CTkFrame(win)
    frm.pack(fill="both", expand=True, padx=16, pady=16)

    ctk.CTkLabel(frm, text="Provider 名称（如 xai / moonshot）", font=("", 12)).pack(anchor="w")
    name_entry = ctk.CTkEntry(frm, height=32)
    name_entry.pack(fill="x", pady=(4, 12))

    ctk.CTkLabel(frm, text="Base URL（可选）", font=("", 12)).pack(anchor="w")
    base_entry = ctk.CTkEntry(frm, height=32)
    base_entry.pack(fill="x", pady=(4, 12))

    def _save():
        p_name = _normalize_provider_name(name_entry.get())
        if not p_name:
            messagebox.showerror("新增失败", "Provider 名称不能为空")
            return
        _ensure_provider_exists(p_name)
        if base_entry.get().strip():
            config["LLM"]["PROVIDERS"][p_name]["BASE_URL"] = base_entry.get().strip()
        save_config(config)
        try:
            on_confirm(p_name)
        finally:
            win.destroy()

    btns = ctk.CTkFrame(frm)
    btns.pack(fill="x", pady=(8, 0))
    ctk.CTkButton(btns, text="取消", width=100, height=32, command=win.destroy).pack(side="right")
    ctk.CTkButton(btns, text="保存", width=100, height=32, fg_color="#4CAF50", text_color="white", command=_save).pack(side="right", padx=10)


def make_mask_widget(name, parent):
    frame = ctk.CTkFrame(parent)
    lbl = ctk.CTkLabel(frame, text="已设置（隐藏）", text_color="gray")
    def on_show():
        reveal_entry(name, parent, frame)
    btn = ctk.CTkButton(frame, text="显示", width=60, height=28, command=on_show)
    lbl.pack(side="left")
    btn.pack(side="left", padx=8)
    return frame


def reveal_entry(name, parent, mask_frame=None):
    if mask_frame:
        mask_frame.destroy()
    ent = ctk.CTkEntry(parent, height=32)
    ent.insert(0, hidden_api_keys.get(name, ""))
    ent.pack(side="left", fill="x", expand=True)
    ent.bind("<FocusOut>", lambda e: on_focus_out_mask(name, ent, parent))
    entries_map[name] = ent


def on_focus_out_mask(name, entry_widget, parent):
    val = entry_widget.get().strip()
    if val:
        hidden_api_keys[name] = val
        entry_widget.destroy()
        mask = make_mask_widget(name, parent)
        mask.pack(side="left")
        entries_map[name] = None


def get_api_key_value(name):
    ent = entries_map.get(name)
    if ent:
        return ent.get().strip()
    return hidden_api_keys.get(name)

# ========== 任务队列相关 ==========
TASK_STATUS_LABELS = {
    "pending": "待完成",
    "running": "正在改",
    "done": "修改完成",
    "failed": "修改失败",
}
TASK_STATUS_TAGS = {
    "pending": "pending",
    "running": "running",
    "done": "done",
    "failed": "failed",
}

task_queue = []

def infer_student_and_essay(folder_name: str) -> tuple[str, str]:
    if "_" in folder_name:
        return tuple(folder_name.split("_", 1))
    if "-" in folder_name:
        return tuple(folder_name.split("-", 1))
    return folder_name, folder_name


def count_existing_docx_chars(folder_path: str) -> str:
    try:
        doc_name = os.path.basename(folder_path)
        docx_path = os.path.join(folder_path, f"{doc_name}.docx")
        if os.path.isfile(docx_path):
            doc = Document(docx_path)
            total = sum(len(p.text.strip()) for p in doc.paragraphs if p.text.strip())
            return str(total)
    except Exception:
        pass
    return ""


def has_images_folder(path: str) -> bool:
    try:
        return any(
            f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))
            for f in os.listdir(path)
            if os.path.isfile(os.path.join(path, f))
        )
    except Exception:
        return False


def scan_folder_for_tasks(folder: str) -> list[str]:
    folder = os.path.abspath(folder)
    tasks = []
    if has_images_folder(folder):
        tasks.append(folder)
    try:
        for name in sorted(os.listdir(folder)):
            if name == "旧":
                continue
            child = os.path.join(folder, name)
            if os.path.isdir(child) and has_images_folder(child):
                tasks.append(child)
    except Exception:
        pass
    return tasks


def render_task_queue():
    queue_tree.delete(*queue_tree.get_children())
    for index, task_path in enumerate(task_queue, start=1):
        folder_name = os.path.basename(task_path)
        student_name, essay_name = infer_student_and_essay(folder_name)
        before_count = count_existing_docx_chars(task_path)
        queue_tree.insert(
            "",
            "end",
            iid=task_path,
            values=(
                str(index),
                student_name,
                task_path,
                essay_name,
                before_count,
                TASK_STATUS_LABELS.get("pending", "待完成"),
            ),
            tags=(TASK_STATUS_TAGS.get("pending", "pending"),),
        )


def refresh_task_queue():
    global task_queue
    task_queue = [p for p in task_queue if os.path.isdir(p)]
    render_task_queue()


def add_task_to_queue(task_path: str) -> bool:
    task_path = os.path.abspath(task_path)
    if not os.path.isdir(task_path) or not has_images_folder(task_path):
        return False
    if task_path in task_queue:
        return False
    task_queue.append(task_path)
    render_task_queue()
    return True


def load_task_queue_from_folder():
    folder = path_entry.get().strip()
    if not folder or not os.path.isdir(folder):
        append_log('❌ 当前路径无效，无法读取任务队列')
        return

    candidates = scan_folder_for_tasks(folder)
    added = 0
    for task_path in candidates:
        if task_path not in task_queue:
            task_queue.append(task_path)
            added += 1
    if added:
        append_log(f'✅ 已读取并加入 {added} 个任务')
    else:
        append_log('ℹ️ 当前路径下没有新任务可加入')
    render_task_queue()


def add_task_button():
    folder = filedialog.askdirectory()
    if not folder:
        return
    if add_task_to_queue(folder):
        append_log(f'✅ 已添加任务：{folder}')
    else:
        append_log(f'❌ 无效文件夹或已存在任务：{folder}')


def remove_selected_task():
    selected = queue_tree.selection()
    if not selected:
        append_log('❌ 请先选择要删除的队列项')
        return
    for item_id in selected:
        if item_id in task_queue:
            task_queue.remove(item_id)
        queue_tree.delete(item_id)
    render_task_queue()


def update_task_status(task_path: str, status: str):
    if queue_tree is None or status not in TASK_STATUS_LABELS:
        return
    if task_path not in queue_tree.get_children():
        return
    item = queue_tree.item(task_path)
    values = list(item["values"])
    if len(values) >= 6:
        values[5] = TASK_STATUS_LABELS[status]
    queue_tree.item(task_path, values=values, tags=(TASK_STATUS_TAGS[status],))

ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

root = ctk.CTk()
root.title("Composition OCR Assistant 作文修改助手 v1.0")

# 固定窗口大小，不受屏幕分辨率影响
win_w, win_h = 1100, 800
root.geometry(f"{win_w}x{win_h}")
root.resizable(True, True)

def _resource_path(name: str) -> Path:
    """兼容源码/打包环境的资源路径解析。"""
    base_dir = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return base_dir / name


def _set_app_icon(win):
    """统一设置窗口左上角与任务栏图标。"""
    try:
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("CompositionOCRAssistant.App")
    except Exception:
        pass

    ico_candidates = [
        _resource_path("app.ico"),
        Path.cwd() / "app.ico",
        Path(__file__).resolve().parent / "app.ico",
    ]

    for ico_path in ico_candidates:
        try:
            if ico_path.exists():
                abs_ico = str(ico_path.resolve())
                # Windows 标题栏图标优先使用 iconbitmap(default=...)
                try:
                    win.iconbitmap(default=abs_ico)
                except TypeError:
                    win.iconbitmap(abs_ico)
                # 某些 Windows 环境里，仅 iconbitmap 不会同步到任务栏，补一次 iconphoto。
                img = Image.open(abs_ico)
                photo = ImageTk.PhotoImage(img)
                win.iconphoto(True, photo)
                win._icon_photo = photo
                # 某些 CustomTkinter 主题切换/初始化会覆盖图标，空闲后再补一次。
                win.after(200, lambda p=abs_ico: _reapply_icon(win, p))
                return
        except Exception:
            continue

    # 回退：使用 base64 内嵌图标，避免完全无图标。
    try:
        icon_data = base64.b64decode(icon_base64.strip())
        image = Image.open(io.BytesIO(icon_data)).resize((32, 32), Image.Resampling.LANCZOS)
        photo = ImageTk.PhotoImage(image)
        win.iconphoto(True, photo)
        win._icon_photo = photo
    except Exception:
        try:
            photo = tk.PhotoImage(data=icon_base64.strip())
            win.iconphoto(True, photo)
            win._icon_photo = photo
        except Exception:
            pass


def _reapply_icon(win, abs_ico: str):
    try:
        win.iconbitmap(default=abs_ico)
    except Exception:
        try:
            win.iconbitmap(abs_ico)
        except Exception:
            pass


_set_app_icon(root)


# ========== 可折叠框架 ==========
class CollapsibleFrame(ctk.CTkFrame):
    """可折叠/展开的框架组件，content_frame 始终保留在布局中，通过隐藏子控件实现收起"""
    def __init__(self, master, title="", collapsed=True, **kwargs):
        super().__init__(master, **kwargs)
        self._collapsed = collapsed

        # 标题栏
        self.title_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.title_frame.pack(fill="x", padx=4, pady=(4, 0))

        self.toggle_btn = ctk.CTkButton(
            self.title_frame,
            text=("▶ " + title) if collapsed else ("▼ " + title),
            width=200,
            height=28,
            anchor="w",
            fg_color="transparent",
            text_color=("#1a73e8", "#8ab4f8"),
            font=("", 13, "bold"),
            hover_color=("#e8f0fe", "#1a2744"),
            command=self._toggle,
        )
        self.toggle_btn.pack(side="left")

        # 内容容器 - 始终在 pack 列表中，不移除
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.content_frame.pack(fill="x", padx=4, pady=(0, 4))

        if collapsed:
            self._hide_children()

    def _toggle(self):
        if self._collapsed:
            self._show_children()
            self._collapsed = False
            self.toggle_btn.configure(text="▼ " + self.toggle_btn.cget("text")[2:])
        else:
            self._hide_children()
            self._collapsed = True
            self.toggle_btn.configure(text="▶ " + self.toggle_btn.cget("text")[2:])

    def _hide_children(self):
        """隐藏 content_frame 内所有子控件"""
        for child in self.content_frame.winfo_children():
            child.pack_forget()

    def _show_children(self):
        """重新显示 content_frame 内所有子控件"""
        for child in self.content_frame.winfo_children():
            child.pack(fill="x", pady=(4, 0))

    def get_content_frame(self):
        return self.content_frame

# ========== 页面切换框架 ==========
pages = {}
current_page = tk.StringVar(value="ocr")

def show_page(page_name):
    """显示指定页面"""
    current_page.set(page_name)
    for page in pages.values():
        page.pack_forget()
    pages[page_name].pack(fill="both", expand=True)

# 顶部按钮框
top_frame = ctk.CTkFrame(root, fg_color=("#f0f0f0", "#2b2b2b"))
top_frame.pack(side="top", fill="x", padx=8, pady=(8, 4))

ctk.CTkLabel(top_frame, text="功能选择:", text_color="gray", font=("", 13, "bold")).pack(side="left", padx=(8, 12))
ctk.CTkButton(top_frame, text="📷 图片转作文", width=140, height=32, command=lambda: show_page("ocr")).pack(side="left", padx=4)
ctk.CTkButton(top_frame, text="📝 docx作文处理", width=140, height=32, command=lambda: show_page("ai")).pack(side="left", padx=4)

# 右上角：配置编辑（打开 JSON 配置的统一编辑窗口）
ctk.CTkButton(top_frame, text="⚙️ 配置编辑", width=120, height=32, command=lambda: open_config_editor()).pack(side="right", padx=8)

# ========== PAGE 1: OCR 处理 ==========
page1 = ctk.CTkScrollableFrame(root)
pages["ocr"] = page1

# ========== 百度图片矫正配置（最前面） - 可折叠 ==========
baidu_collapse = CollapsibleFrame(page1, title="📷 百度图片矫正（OCR前自动矫正倾斜/弯曲文档）", collapsed=True)
baidu_collapse.pack(padx=12, fill="x", pady=(8, 0))
baidu_content = baidu_collapse.get_content_frame()

# 启用图片矫正开关
baidu_enable_frame = ctk.CTkFrame(baidu_content, fg_color="transparent")
baidu_enable_frame.pack(fill="x", pady=(4, 0))
use_baidu_correction_var = tk.BooleanVar(value=bool(config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("ENABLED", False)))
ctk.CTkCheckBox(baidu_enable_frame, text="启用图片矫正（去阴影+透视变换）", variable=use_baidu_correction_var).pack(side="left", padx=4)

# 百度 API Key
baidu_key_frame = ctk.CTkFrame(baidu_content, fg_color="transparent")
baidu_key_frame.pack(fill="x", pady=(4, 0))
ctk.CTkLabel(baidu_key_frame, text="百度 API Key", width=110, anchor="w").pack(side="left", padx=(0, 8))
baidu_api_key_entry = ctk.CTkEntry(baidu_key_frame)
baidu_api_key_entry.insert(0, config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("API_KEY", ""))
baidu_api_key_entry.pack(side="left", fill="x", expand=True)
entries_map['baidu_api_key'] = baidu_api_key_entry
if config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("API_KEY"):
    hidden_api_keys['baidu_api_key'] = config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("API_KEY", "")
    baidu_api_key_entry.destroy()
    mask = make_mask_widget('baidu_api_key', baidu_key_frame)
    mask.pack(side="left")
    entries_map['baidu_api_key'] = None

# 百度 Secret Key
baidu_secret_frame = ctk.CTkFrame(baidu_content, fg_color="transparent")
baidu_secret_frame.pack(fill="x", pady=(4, 0))
ctk.CTkLabel(baidu_secret_frame, text="百度 Secret Key", width=110, anchor="w").pack(side="left", padx=(0, 8))
baidu_secret_key_entry = ctk.CTkEntry(baidu_secret_frame)
baidu_secret_key_entry.insert(0, config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY", ""))
baidu_secret_key_entry.pack(side="left", fill="x", expand=True)
entries_map['baidu_secret_key'] = baidu_secret_key_entry
if config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY"):
    hidden_api_keys['baidu_secret_key'] = config.get("OCR", {}).get("BAIDU_CORRECTION", {}).get("SECRET_KEY", "")
    baidu_secret_key_entry.destroy()
    mask = make_mask_widget('baidu_secret_key', baidu_secret_frame)
    mask.pack(side="left")
    entries_map['baidu_secret_key'] = None

# OCR 配置 - 可折叠
ocr_collapse = CollapsibleFrame(page1, title="🔍 OCR 识别配置", collapsed=True)
ocr_collapse.pack(padx=12, fill="x", pady=(8, 0))
ocr_content = ocr_collapse.get_content_frame()

# OCR 配置（标签与输入框同一行）
ocr_frame = ctk.CTkFrame(ocr_content, fg_color="transparent")
ocr_frame.pack(fill="x", pady=(4, 0))
ctk.CTkLabel(ocr_frame, text="OCR 接口 URL", width=110, anchor="w").pack(side="left", padx=(0, 8))
url_entry = ctk.CTkEntry(ocr_frame)
url_entry.insert(0, config.get("OCR", {}).get("XFYUN", {}).get("URL", ""))
url_entry.pack(side="left", fill="x", expand=True)

# APPID 与 API_KEY 同行布局
appid_frame = ctk.CTkFrame(ocr_content, fg_color="transparent")
appid_frame.pack(fill="x", pady=(4, 0))
ctk.CTkLabel(appid_frame, text="APPID", width=110, anchor="w").pack(side="left", padx=(0, 8))
appid_entry = ctk.CTkEntry(appid_frame)
appid_entry.insert(0, config.get("OCR", {}).get("XFYUN", {}).get("APPID", ""))
appid_entry.pack(side="left", fill="x", expand=True)

apikey_frame = ctk.CTkFrame(ocr_content, fg_color="transparent")
apikey_frame.pack(fill="x", pady=(4, 0))
ctk.CTkLabel(apikey_frame, text="API_KEY", width=110, anchor="w").pack(side="left", padx=(0, 8))
apikey_entry = ctk.CTkEntry(apikey_frame)
apikey_entry.insert(0, config.get("OCR", {}).get("XFYUN", {}).get("API_KEY", ""))
apikey_entry.pack(side="left", fill="x", expand=True)
entries_map['ocr'] = apikey_entry
if config.get("OCR", {}).get("XFYUN", {}).get("API_KEY"):
    hidden_api_keys['ocr'] = config.get("OCR", {}).get("XFYUN", {}).get("API_KEY", "")
    apikey_entry.destroy()
    mask = make_mask_widget('ocr', apikey_frame)
    mask.pack(side="left")
    entries_map['ocr'] = None

# 第一步 AI 改错别字 - 分区标题
typo_task_provider = _normalize_provider_name(
    (config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROVIDER", "deepseek")
) or "deepseek"
editor_task_provider = _normalize_provider_name(
    (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")
) or "deepseek"
_ensure_provider_exists(typo_task_provider)
_ensure_provider_exists(editor_task_provider)
provider_names = _provider_name_list()

ai_typo_title = ctk.CTkFrame(page1, fg_color=("#e6f4ea", "#1a3326"))
ai_typo_title.pack(padx=12, fill="x", pady=(12, 0))
ctk.CTkLabel(ai_typo_title, text="🧠 第一步：AI 错别字修正", font=("", 13, "bold"), text_color=("#137333", "#81c995")).pack(side="left", anchor="w", padx=8, pady=6)

deepseek_frame = ctk.CTkFrame(page1)
deepseek_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(deepseek_frame, text="AI API Key", width=110, anchor="w").pack(side="left", padx=(0, 8))
deepseek_entry = ctk.CTkEntry(deepseek_frame)
deepseek_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(typo_task_provider, {}) or {}).get("API_KEY", ""))
deepseek_entry.pack(side="left", fill="x", expand=True)
entries_map['deepseek'] = deepseek_entry
# 如果已配置 provider 的 API_KEY，则默认隐藏显示
if (config.get("LLM", {}).get("PROVIDERS", {}).get(typo_task_provider, {}) or {}).get("API_KEY"):
    hidden_api_keys['deepseek'] = (config.get("LLM", {}).get("PROVIDERS", {}).get(typo_task_provider, {}) or {}).get("API_KEY", "")
    deepseek_entry.destroy()
    mask = make_mask_widget('deepseek', deepseek_frame)
    mask.pack(side="left")
    entries_map['deepseek'] = None

use_deepseek_var = tk.BooleanVar(value=bool((config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("ENABLED", False)))
ctk.CTkCheckBox(deepseek_frame, text="启用 AI 错别字自动修正（较慢）", variable=use_deepseek_var).pack(side="left", padx=8)

# DeepSeek Base URL 同行
deepseek_base_frame = ctk.CTkFrame(page1)
deepseek_base_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(deepseek_base_frame, text="AI Provider", width=110, anchor="w").pack(side="left", padx=(0, 8))
deepseek_provider_name = typo_task_provider
deepseek_provider_menu = ctk.CTkOptionMenu(deepseek_base_frame, values=provider_names, width=140)
deepseek_provider_menu.set(deepseek_provider_name)
deepseek_provider_menu.pack(side="left", padx=(0, 8))
ctk.CTkButton(deepseek_base_frame, text="+ 新增", width=68, command=lambda: _open_new_provider_dialog(_on_new_typo_provider)).pack(side="left", padx=(0, 8))

ctk.CTkLabel(deepseek_base_frame, text="Base URL", width=70, anchor="w").pack(side="left", padx=(8, 4))
deepseek_base_entry = ctk.CTkEntry(deepseek_base_frame)
deepseek_base_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(deepseek_provider_name, {}) or {}).get("BASE_URL", ""))
deepseek_base_entry.pack(side="left", fill="x", expand=True)

def _on_deepseek_provider_change(choice):
    p_name = _ensure_provider_exists(choice)
    deepseek_base_entry.configure(state="normal")
    deepseek_base_entry.delete(0, tk.END)
    deepseek_base_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

def _on_new_typo_provider(p_name):
    names = _provider_name_list()
    deepseek_provider_menu.configure(values=names)
    deepseek_provider_menu.set(p_name)
    _on_deepseek_provider_change(p_name)

deepseek_provider_menu.configure(command=_on_deepseek_provider_change)
_on_deepseek_provider_change(deepseek_provider_name)

# Prompt（多行）
prompt_frame = ctk.CTkFrame(page1)
prompt_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(prompt_frame, text="自定义提示词", width=110, anchor="nw").pack(side="left", padx=(0, 8), pady=(4, 0))
prompt_text = ctk.CTkTextbox(prompt_frame, height=120)
prompt_text.insert(
    "1.0",
    (config.get("LLM", {}).get("TASKS", {}).get("typo_fix", {}) or {}).get("PROMPT")
    or DEFAULT_CONFIG["LLM"]["TASKS"]["typo_fix"]["PROMPT"],
)
prompt_text.pack(side="left", fill="x", expand=True)

# 编辑 API（第二步）- 分区标题
editor_title = ctk.CTkFrame(page1, fg_color=("#fce8e6", "#3c1a1a"))
editor_title.pack(padx=12, fill="x", pady=(12, 0))
ctk.CTkLabel(editor_title, text="✍️ 第二步：AI 修改作文", font=("", 13, "bold"), text_color=("#c5221f", "#f28b82")).pack(side="left", anchor="w", padx=8, pady=6)

use_editor_var = tk.BooleanVar(value=bool((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("ENABLED", False)))
editor_enable_frame = ctk.CTkFrame(page1)
editor_enable_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkCheckBox(editor_enable_frame, text="启用 第二步 修改作文", variable=use_editor_var).pack(side="left", padx=4)

# 第二步 API Key 同行
editor_key_frame = ctk.CTkFrame(page1)
editor_key_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(editor_key_frame, text="AI API Key", width=110, anchor="w").pack(side="left", padx=(0, 8))
editor_key_entry = ctk.CTkEntry(editor_key_frame)
editor_key_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(editor_task_provider, {}) or {}).get("API_KEY", ""))
editor_key_entry.pack(side="left", fill="x", expand=True)
entries_map['editor'] = editor_key_entry
if (config.get("LLM", {}).get("PROVIDERS", {}).get(editor_task_provider, {}) or {}).get("API_KEY"):
    hidden_api_keys['editor'] = (config.get("LLM", {}).get("PROVIDERS", {}).get(editor_task_provider, {}) or {}).get("API_KEY", "")
    editor_key_entry.destroy()
    mask = make_mask_widget('editor', editor_key_frame)
    mask.pack(side="left")
    entries_map['editor'] = None

# 第二步 Base URL 同行
editor_base_frame = ctk.CTkFrame(page1)
editor_base_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(editor_base_frame, text="AI Provider", width=110, anchor="w").pack(side="left", padx=(0, 8))
editor_provider_name = editor_task_provider
editor_provider_menu = ctk.CTkOptionMenu(editor_base_frame, values=provider_names, width=140)
editor_provider_menu.set(editor_provider_name)
editor_provider_menu.pack(side="left", padx=(0, 8))
ctk.CTkButton(editor_base_frame, text="+ 新增", width=68, command=lambda: _open_new_provider_dialog(_on_new_editor_provider)).pack(side="left", padx=(0, 8))

ctk.CTkLabel(editor_base_frame, text="Base URL", width=70, anchor="w").pack(side="left", padx=(8, 4))
editor_base_entry = ctk.CTkEntry(editor_base_frame)
editor_base_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(editor_provider_name, {}) or {}).get("BASE_URL", ""))
editor_base_entry.pack(side="left", fill="x", expand=True)

def _on_editor_provider_change(choice):
    p_name = _ensure_provider_exists(choice)
    editor_base_entry.configure(state="normal")
    editor_base_entry.delete(0, tk.END)
    editor_base_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {}).get("BASE_URL", ""))

def _on_new_editor_provider(p_name):
    names = _provider_name_list()
    editor_provider_menu.configure(values=names)
    editor_provider_menu.set(p_name)
    _on_editor_provider_change(p_name)

editor_provider_menu.configure(command=_on_editor_provider_change)
_on_editor_provider_change(editor_provider_name)

# 第二步 Prompt 多行
editor_prompt_frame = ctk.CTkFrame(page1)
editor_prompt_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(editor_prompt_frame, text="自定义提示词", width=110, anchor="nw").pack(side="left", padx=(0, 8), pady=(4, 0))
editor_prompt_text = ctk.CTkTextbox(editor_prompt_frame, height=120)
editor_prompt_text.insert(
    "1.0",
    (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT")
    or DEFAULT_CONFIG["LLM"]["TASKS"]["editor"]["PROMPT"],
)
editor_prompt_text.pack(side="left", fill="x", expand=True)

# 目标字数范围
editor_word_count_frame = ctk.CTkFrame(page1)
editor_word_count_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(editor_word_count_frame, text="目标字数", width=110, anchor="w").pack(side="left", padx=(0, 8))
editor_count_min_entry = ctk.CTkEntry(editor_word_count_frame, width=100)
editor_count_min_entry.insert(
    0,
    str((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""),
)
editor_count_min_entry.pack(side="left", padx=(0, 4))
ctk.CTkLabel(editor_word_count_frame, text="-", width=10).pack(side="left")
editor_count_max_entry = ctk.CTkEntry(editor_word_count_frame, width=100)
editor_count_max_entry.insert(
    0,
    str((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""),
)
editor_count_max_entry.pack(side="left", padx=(4, 0))
ctk.CTkLabel(editor_word_count_frame, text="（空白表示自动）", text_color="gray").pack(side="left", padx=(8, 0))

# 路径 - 分区标题
path_title = ctk.CTkFrame(page1, fg_color=("#fef7e0", "#3c341a"))
path_title.pack(padx=12, fill="x", pady=(12, 0))
ctk.CTkLabel(path_title, text="📁 文件路径与任务", font=("", 13, "bold"), text_color=("#b06000", "#fdd663")).pack(side="left", anchor="w", padx=8, pady=6)

path_frame = ctk.CTkFrame(page1)
path_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(path_frame, text="作文文件夹路径", width=110, anchor="w").pack(side="left", padx=(0, 8))
path_entry = ctk.CTkEntry(path_frame)
path_entry.insert(0, config["APP"]["ROOT_DIR"])
path_entry.pack(side="left", fill="x", expand=True)
ctk.CTkButton(path_frame, text="浏览", width=70, command=browse_folder).pack(side="left", padx=8)

# 任务队列
queue_title_frame = ctk.CTkFrame(page1)
queue_title_frame.pack(padx=12, fill="x", pady=(10, 0))
ctk.CTkLabel(queue_title_frame, text="任务队列", font=("", 13, "bold")).pack(side="left", anchor="w")

queue_control_frame = ctk.CTkFrame(page1)
queue_control_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkButton(queue_control_frame, text="➕ 添加", width=90, command=add_task_button).pack(side="left")
ctk.CTkButton(queue_control_frame, text="🗑️ 删除", width=90, command=remove_selected_task).pack(side="left", padx=8)
ctk.CTkButton(queue_control_frame, text="📥 读取", width=90, command=load_task_queue_from_folder).pack(side="right")
ctk.CTkButton(queue_control_frame, text="🔄 刷新", width=90, command=refresh_task_queue).pack(side="right", padx=(0, 8))

queue_frame = ctk.CTkFrame(page1)
queue_frame.pack(padx=12, fill="both", pady=(6, 0), expand=False)

queue_scroll_y = tk.Scrollbar(queue_frame, orient="vertical")
queue_scroll_y.pack(side="right", fill="y")
queue_scroll_x = tk.Scrollbar(queue_frame, orient="horizontal")
queue_scroll_x.pack(side="bottom", fill="x")

queue_tree = ttk.Treeview(
    queue_frame,
    columns=("index", "student", "path", "essay", "count", "status"),
    show="headings",
    yscrollcommand=queue_scroll_y.set,
    xscrollcommand=queue_scroll_x.set,
    selectmode="browse",
)
queue_tree.heading("index", text="序号")
queue_tree.heading("student", text="学生姓名")
queue_tree.heading("path", text="文件路径")
queue_tree.heading("essay", text="作文名称")
queue_tree.heading("count", text="修改前字数")
queue_tree.heading("status", text="任务状态")
queue_tree.column("index", width=50, anchor="center", minwidth=40)
queue_tree.column("student", width=120, anchor="w", minwidth=80)
queue_tree.column("path", width=300, anchor="w", minwidth=150)
queue_tree.column("essay", width=150, anchor="w", minwidth=80)
queue_tree.column("count", width=100, anchor="center", minwidth=70)
queue_tree.column("status", width=100, anchor="center", minwidth=80)
queue_tree.pack(side="left", fill="both", expand=True)
queue_scroll_y.config(command=queue_tree.yview)
queue_scroll_x.config(command=queue_tree.xview)

style = ttk.Style()
style.configure("Treeview", rowheight=32, font=("", 12))
style.configure("Treeview.Heading", font=("", 12, "bold"))
queue_tree.tag_configure("pending", background="#cfe2ff")
queue_tree.tag_configure("running", background="#fff3bf")
queue_tree.tag_configure("done", background="#d4edda")
queue_tree.tag_configure("failed", background="#f8d7da")
refresh_task_queue()

# 启动按钮
ctk.CTkButton(page1, text="🚀 开始处理", fg_color="#4CAF50", text_color="white", height=40, command=start_processing).pack(pady=12)

# 日志
ctk.CTkLabel(page1, text="运行日志", anchor="w").pack(anchor="w", padx=12, pady=(0, 4))
log_text = ctk.CTkTextbox(page1, height=180)
log_text.configure(state="disabled")
log_text.pack(padx=12, pady=(0, 10), fill="both", expand=True)

# ========== PAGE 2: AI 综合处理 - 可拖动任务流 ==========
page2 = ctk.CTkScrollableFrame(root)
pages["ai"] = page2

# 任务列表配置（顺序：6→1→AI→2→3→5）
task_config = [
    {"id": "6", "name": "6. 转换 DOC → DOCX", "enabled": True, "order": 0},
    {"id": "1", "name": "1. 清除空格", "enabled": True, "order": 1},
    {"id": "AI", "name": "🤖 AI 改作文", "enabled": True, "order": 2},
    {"id": "2", "name": '2. 添加"修改前/后"', "enabled": True, "order": 3},
    {"id": "3", "name": "3. 格式化字体段落", "enabled": True, "order": 4},
    {"id": "5", "name": "5. 修改作者", "enabled": True, "order": 5},
]

# AI API Key
ai_task_provider = _normalize_provider_name(
    (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROVIDER", "deepseek")
) or "deepseek"
_ensure_provider_exists(ai_task_provider)

# 分区标题 - AI配置
ai_config_title = ctk.CTkFrame(page2, fg_color=("#e8f0fe", "#1a2744"))
ai_config_title.pack(padx=12, fill="x", pady=(12, 0))
ctk.CTkLabel(ai_config_title, text="🤖 AI 配置", font=("", 13, "bold"), text_color=("#1a73e8", "#8ab4f8")).pack(side="left", anchor="w", padx=8, pady=6)

# 文件夹选择
ai_path_frame = ctk.CTkFrame(page2)
ai_path_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_path_frame, text="处理文件夹", width=110, anchor="w").pack(side="left", padx=(0, 8))
ai_path_entry = ctk.CTkEntry(ai_path_frame)
ai_path_entry.pack(side="left", fill="x", expand=True)
def browse_ai_folder():
    folder = filedialog.askdirectory()
    if folder:
        ai_path_entry.delete(0, tk.END)
        ai_path_entry.insert(0, folder)
ctk.CTkButton(ai_path_frame, text="浏览", command=browse_ai_folder, width=70).pack(side="left", padx=8)

ai_provider_frame = ctk.CTkFrame(page2)
ai_provider_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_provider_frame, text="AI Provider", width=110, anchor="w").pack(side="left", padx=(0, 8))
ai_provider_menu = ctk.CTkOptionMenu(ai_provider_frame, values=_provider_name_list(), width=140)
ai_provider_menu.set(ai_task_provider)
ai_provider_menu.pack(side="left", padx=(0, 8))

def _on_ai_provider_change(choice):
    p_name = _ensure_provider_exists(choice)
    provider_cfg = (config.get("LLM", {}).get("PROVIDERS", {}).get(p_name, {}) or {})
    ai_key_entry.delete(0, tk.END)
    ai_key_entry.insert(0, provider_cfg.get("API_KEY", ""))
    ai_url_entry.delete(0, tk.END)
    ai_url_entry.insert(0, provider_cfg.get("BASE_URL", ""))

def _on_new_ai_provider(p_name):
    ai_provider_menu.configure(values=_provider_name_list())
    ai_provider_menu.set(p_name)
    _on_ai_provider_change(p_name)

ctk.CTkButton(ai_provider_frame, text="+ 新增", width=68, command=lambda: _open_new_provider_dialog(_on_new_ai_provider)).pack(side="left")
ai_provider_menu.configure(command=_on_ai_provider_change)

ai_key_frame = ctk.CTkFrame(page2)
ai_key_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_key_frame, text="AI API Key", width=110, anchor="w").pack(side="left", padx=(0, 8))
ai_key_entry = ctk.CTkEntry(ai_key_frame, show="*")
ai_key_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(ai_task_provider, {}) or {}).get("API_KEY", ""))
ai_key_entry.pack(side="left", fill="x", expand=True)

# Base URL
ai_url_frame = ctk.CTkFrame(page2)
ai_url_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_url_frame, text="Base URL", width=110, anchor="w").pack(side="left", padx=(0, 8))
ai_url_entry = ctk.CTkEntry(ai_url_frame)
ai_url_entry.insert(0, (config.get("LLM", {}).get("PROVIDERS", {}).get(ai_task_provider, {}) or {}).get("BASE_URL", ""))
ai_url_entry.pack(side="left", fill="x", expand=True)

# Prompt
ai_prompt_frame = ctk.CTkFrame(page2)
ai_prompt_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_prompt_frame, text="AI Prompt", width=110, anchor="nw").pack(side="left", padx=(0, 8), pady=(4, 0))
ai_prompt_text = ctk.CTkTextbox(ai_prompt_frame, height=80)
ai_prompt_text.insert(
    "1.0",
    (config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("PROMPT", "{text}"),
)
ai_prompt_text.pack(side="left", fill="x", expand=True)

# 手动字数范围
ai_word_count_frame = ctk.CTkFrame(page2)
ai_word_count_frame.pack(padx=12, fill="x", pady=(6, 0))
ctk.CTkLabel(ai_word_count_frame, text="目标字数", width=110, anchor="w").pack(side="left", padx=(0, 8))
ai_count_min_entry = ctk.CTkEntry(ai_word_count_frame, width=100)
ai_count_min_entry.insert(
    0,
    str((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MIN") or ""),
)
ai_count_min_entry.pack(side="left", padx=(0, 4))
ctk.CTkLabel(ai_word_count_frame, text="-", width=10).pack(side="left")
ai_count_max_entry = ctk.CTkEntry(ai_word_count_frame, width=100)
ai_count_max_entry.insert(
    0,
    str((config.get("LLM", {}).get("TASKS", {}).get("editor", {}) or {}).get("COUNT_MAX") or ""),
)
ai_count_max_entry.pack(side="left", padx=(4, 0))
ctk.CTkLabel(ai_word_count_frame, text="（空白表示自动）", text_color="gray").pack(side="left", padx=(8, 0))

# 任务流标签框架 - 可拖动和排序
task_flow_title = ctk.CTkFrame(page2, fg_color=("#e6f4ea", "#1a3326"))
task_flow_title.pack(padx=12, fill="x", pady=(12, 0))
ctk.CTkLabel(task_flow_title, text="📋 处理流程（勾选/取消步骤，拖动上下移动顺序）", font=("", 13, "bold"), text_color=("#137333", "#81c995")).pack(side="left", anchor="w", padx=8, pady=6)

# 创建任务流程容器框
task_container = ctk.CTkFrame(page2, fg_color=("#ffffff", "#2a2a2a"), border_width=2, border_color=("gray50", "gray50"))
task_container.pack(padx=12, pady=(6, 10), fill="both", expand=False)

task_frames_map = {}  # id -> {"frame": frame, "var": var}

def create_task_frame(task):
    """创建单个任务框，支持拖动排序 - 改进版实现"""
    frame = ctk.CTkFrame(task_container, fg_color=("#f8f9fa", "#1f1f1f"), border_width=1, border_color=("gray70", "gray30"))
    frame.pack(padx=8, fill="x", pady=3)
    
    # 左侧：勾选框 + 拖动把手 + 任务名
    left_frame = ctk.CTkFrame(frame, fg_color="transparent")
    left_frame.pack(side="left", padx=8, pady=6, fill="x", expand=True)
    
    var = tk.BooleanVar(value=task["enabled"])
    # 使用 trace 绑定，确保变量变化总是同步到 task 配置
    def _on_var_changed(*_):
        task["enabled"] = bool(var.get())
    var.trace_add("write", _on_var_changed)
    
    # 拖动把手（只在这个 widget 上绑定拖动事件）
    handle_label = ctk.CTkLabel(left_frame, text="☰", text_color="gray", font=("", 14), cursor="hand2")
    handle_label.pack(side="left", padx=(0, 5))
    
    # 拖动状态存储在 frame 的自定义属性中
    frame._drag_state = {
        "start_y": 0,
        "is_dragging": False,
        "initial_color": ("#f0f0f0", "#1f1f1f"),
        "initial_border": ("gray70", "gray30")
    }
    
    # 复选框（使用 trace 处理变量变化，无需额外回调）
    checkbox = ctk.CTkCheckBox(left_frame, text=task["name"], variable=var)
    checkbox.pack(side="left", fill="x", expand=True)
    
    # ========== 拖动事件处理（只绑定在 handle_label） ==========
    def on_handle_press(event):
        """鼠标按下：记录起始位置"""
        frame._drag_state["start_y"] = event.widget.winfo_rooty() + event.y
        frame._drag_state["is_dragging"] = False
        # 视觉提示：改变颜色表示可拖动
        frame.configure(border_color=("cyan", "lightcyan"), fg_color=("#e8f4f8", "#2a3a3a"))
        # 插入一个占位符（用于吸附效果），并暂时将被拖动的 frame 移除
        try:
            placeholder = ctk.CTkFrame(task_container, height=8, fg_color=("gray80", "gray20"))
            # 在当前 frame 前插入占位符以保持原位
            placeholder.pack(padx=10, fill="x", pady=4, before=frame)
            frame._drag_state["placeholder"] = placeholder
            # 把真实 frame 从布局中移除以便拖动并让占位符可见
            frame.pack_forget()
            frame.lift()
        except Exception:
            frame._drag_state.pop("placeholder", None)
    
    def on_handle_drag(event):
        """鼠标移动：检测并执行拖动"""
        current_y = event.widget.winfo_rooty() + event.y
        delta_y = current_y - frame._drag_state["start_y"]

        # 超过阈值才进入拖动状态
        if abs(delta_y) > 30 and not frame._drag_state["is_dragging"]:
            frame._drag_state["is_dragging"] = True
            frame.configure(border_color=("blue", "lightblue"), fg_color=("#d0e8ff", "#1a2a3a"))

        if not frame._drag_state.get("is_dragging"):
            return

        # 计算当前应该插入的位置（按容器内其它任务的中点比较）
        siblings = [f for f in task_container.winfo_children() if isinstance(f, ctk.CTkFrame) and getattr(f, "_drag_state", None) is not None and f != frame]
        if not siblings:
            frame._drag_state["insert_index"] = 0
            return

        insert_index = len(siblings)
        current_y_abs = event.widget.winfo_rooty() + event.y
        for i, other in enumerate(siblings):
            other_mid = other.winfo_rooty() + other.winfo_height() // 2
            if current_y_abs < other_mid:
                insert_index = i
                break

        # 移动占位符到 insert_index 位置（占位符负责视觉吸附）
        placeholder = frame._drag_state.get("placeholder")
        if placeholder:
            try:
                placeholder.pack_forget()
                if insert_index >= len(siblings):
                    last = siblings[-1]
                    placeholder.pack(padx=10, fill="x", pady=4, after=last)
                else:
                    placeholder.pack(padx=10, fill="x", pady=4, before=siblings[insert_index])
                frame._drag_state["insert_index"] = insert_index
            except Exception:
                pass
    
    def on_handle_release(event):
        """鼠标释放：恢复状态"""
        # 恢复原始视觉状态
        frame.configure(
            border_color=frame._drag_state["initial_border"],
            fg_color=frame._drag_state["initial_color"]
        )
        frame._drag_state["is_dragging"] = False

        # 将真实 frame 插入到占位符所在位置（占位符是在容器中的临时占位）
        insert_index = frame._drag_state.get("insert_index")
        placeholder = frame._drag_state.get("placeholder")

        # 清除占位符并在其原位置插入 frame
        try:
            # 列出当前（含占位符）容器子组件顺序，寻找 placeholder 的位置
            children = list(task_container.winfo_children())
            if placeholder in children:
                idx = children.index(placeholder)
            else:
                idx = None
            # 移除占位符视图
            if placeholder:
                placeholder.pack_forget()
        except Exception:
            idx = None

        # 重新获取当前容器内的任务 frames（不包含正在拖动的 frame）
        frames_now = [f for f in task_container.winfo_children() if isinstance(f, ctk.CTkFrame) and getattr(f, "_drag_state", None) is not None and f != frame]

        # 如果没有确定索引，使用 insert_index（拖动计算）作为回退
        if idx is None:
            if insert_index is None or insert_index >= len(frames_now):
                frame.pack(padx=10, fill="x", pady=4)
            else:
                frame.pack(padx=10, fill="x", pady=4, before=frames_now[insert_index])
        else:
            # 按 insert_index 插入
            if insert_index is None or insert_index >= len(frames_now):
                frame.pack(padx=10, fill="x", pady=4)
            else:
                frame.pack(padx=10, fill="x", pady=4, before=frames_now[insert_index])

        # 最终一次性更新所有任务的 order（仅容器内的任务）
        all_frames_current = [f for f in task_container.winfo_children() if isinstance(f, ctk.CTkFrame) and hasattr(f, "_drag_state")]
        for i, f in enumerate(all_frames_current):
            for task_item in task_config:
                if task_frames_map.get(task_item["id"], {}).get("frame") == f:
                    task_item["order"] = i

        # 清理拖动状态
        frame._drag_state.pop("placeholder", None)
        frame._drag_state.pop("insert_index", None)
    
    # 只在拖动把手上绑定事件
    handle_label.bind("<Button-1>", on_handle_press)
    handle_label.bind("<B1-Motion>", on_handle_drag)
    handle_label.bind("<ButtonRelease-1>", on_handle_release)
    
    # 存储任务框信息
    task_frames_map[task["id"]] = {
        "frame": frame,
        "var": var
    }
    
    return frame

# 创建所有任务框
for task in sorted(task_config, key=lambda x: x["order"]):
    create_task_frame(task)

# 启动按钮
def start_ai_workflow():
    def task():
        folder = ai_path_entry.get().strip()
        selected_provider = _ensure_provider_exists(ai_provider_menu.get() or "deepseek")
        api_key = ai_key_entry.get().strip()
        base_url = ai_url_entry.get().strip()
        prompt = ai_prompt_text.get("1.0", tk.END).strip() or None
        min_text = ai_count_min_entry.get().strip()
        max_text = ai_count_max_entry.get().strip()
        count_min = None
        count_max = None
        if min_text:
            try:
                count_min = int(min_text)
            except ValueError:
                append_log_ai("❌ 目标字数范围最小值必须是整数")
                return
        if max_text:
            try:
                count_max = int(max_text)
            except ValueError:
                append_log_ai("❌ 目标字数范围最大值必须是整数")
                return
        if count_min is not None and count_max is not None and count_min > count_max:
            append_log_ai("❌ 目标字数范围最小值不能大于最大值")
            return
        
        if not folder or not api_key:
            append_log_ai("❌ 请填写文件夹路径和 API Key")
            return
        
        if not os.path.isdir(folder):
            append_log_ai("❌ 文件夹路径无效")
            return
        
        # 保存配置（写入用户选择的 provider）
        config.setdefault("LLM", {})
        config["LLM"].setdefault("PROVIDERS", {})
        config["LLM"].setdefault("TASKS", {})

        config["LLM"]["PROVIDERS"].setdefault(selected_provider, {})
        config["LLM"]["PROVIDERS"][selected_provider]["API_KEY"] = api_key
        config["LLM"]["PROVIDERS"][selected_provider]["BASE_URL"] = base_url
        config["LLM"]["PROVIDERS"][selected_provider].setdefault("MODEL", "deepseek-chat" if selected_provider == "deepseek" else "gpt-4o-mini")

        config["LLM"]["TASKS"].setdefault("editor", {})
        config["LLM"]["TASKS"]["editor"]["PROMPT"] = prompt or "{text}"
        config["LLM"]["TASKS"]["editor"]["ENABLED"] = True
        config["LLM"]["TASKS"]["editor"]["PROVIDER"] = selected_provider
        config["LLM"]["TASKS"]["editor"]["COUNT_MIN"] = count_min
        config["LLM"]["TASKS"]["editor"]["COUNT_MAX"] = count_max

        save_config(config)
        
        append_log_ai("📋 开始处理流程...")
        
        # 第一步：复制所有原始文件为"改 XXX"（包含 docx 与常见图片格式，最多递归 4 层）
        append_log_ai("【准备】复制原始文件（包含图片）...")
        import shutil
        copied_files = []
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".gif")
        for root, files in iter_files_limited(folder, max_depth=4):
            for file in files:
                name_check = file.lstrip()
                name_lower = file.lower()
                if (name_lower.endswith(".docx") or name_lower.endswith(image_exts)) and not name_check.startswith("~$") and not name_check.startswith("改 "):
                    original_path = os.path.join(root, file)
                    new_filename = f"改 {file}"
                    new_path = os.path.join(root, new_filename)
                    try:
                        shutil.copy2(original_path, new_path)
                        copied_files.append(new_filename)
                        append_log_ai(f"  ✓ {new_filename}")
                    except Exception as e:
                        append_log_ai(f"  ✗ {file} 复制失败: {e}")
        
        if not copied_files:
            append_log_ai("❌ 未找到需要处理的文件")
            return
        
        # 获取启用的任务列表
        enabled_tasks = sorted(
            [(t["id"], t["order"]) for t in task_config if t["enabled"]],
            key=lambda x: x[1]
        )
        
        try:
            for task_id, _ in enabled_tasks:
                if task_id == "6":
                    append_log_ai("【步骤 6】转换 DOC → DOCX...")
                    _convert_docs(folder, process_only_modified=True)
                elif task_id == "1":
                    append_log_ai("【步骤 1】清除空格...")
                    _clear_spaces(folder, process_only_modified=True)
                elif task_id == "AI":
                    append_log_ai("【步骤 AI】发送给 AI 修正...")
                    _process_ai(
                        folder,
                        api_key,
                        base_url,
                        prompt,
                        process_only_modified=True,
                        count_min=count_min,
                        count_max=count_max,
                    )
                elif task_id == "2":
                    append_log_ai("【步骤 2】添加标签...")
                    _add_labels(folder, process_only_modified=True)
                elif task_id == "3":
                    append_log_ai("【步骤 3】格式化...")
                    _format_docs(folder, process_only_modified=True)
                elif task_id == "5":
                    append_log_ai("【步骤 5】修改作者...")
                    _set_author(folder, process_only_modified=True)
            
            append_log_ai("✅ 所有流程完成！")
        except Exception as e:
            append_log_ai(f"❌ 处理失败：{e}")
            import traceback
            traceback.print_exc()
    
    threading.Thread(target=task, daemon=True).start()

def append_log_ai(message: str):
    timestamp = datetime.now().strftime("%H:%M:%S")
    ai_log_text.configure(state="normal")
    ai_log_text.insert(tk.END, f"[{timestamp}] {message}\n")
    ai_log_text.see(tk.END)
    ai_log_text.configure(state="disabled")

# 各步骤的处理函数
def _convert_docs(folder, process_only_modified=False):
    """步骤 6：转换 DOC → DOCX"""
    import subprocess
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            # 兼容文件名前导空格与中文名
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".doc") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    cmd = ["soffice", "--headless", "--convert-to", "docx", doc_path, "--outdir", root]
                    subprocess.run(cmd, capture_output=True, timeout=30)
                    base_name = os.path.splitext(os.path.basename(doc_path))[0]
                    new_path = os.path.join(root, base_name + ".docx")
                    if os.path.exists(new_path):
                        os.remove(doc_path)
                        append_log_ai(f"  ✓ {base_name}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")

def _clear_spaces(folder, process_only_modified=False):
    """步骤 1：清除空格"""
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    for para in doc.paragraphs:
                        for run in para.runs:
                            run.text = run.text.strip()
                    doc.save(doc_path)
                    append_log_ai(f"  ✓ {file}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")


def count_chinese_characters(text: str) -> int:
    return sum(1 for ch in text if not ch.isspace())


def determine_word_count_bounds(original_count: int) -> tuple[int, int]:
    if original_count >= 850:
        return max(700, original_count - 30), original_count + 30
    if original_count >= 800:
        return 820, 850
    return 700, 820


def is_word_count_within_rules(text: str, original_count: int) -> bool:
    count = count_chinese_characters(text)
    min_count, max_count = determine_word_count_bounds(original_count)
    return min_count <= count <= max_count


def llm_edit_with_word_count_supervision(
    prompt_template: str,
    full_text: str,
    client: OpenAI,
    model: str,
    original_count: int,
    file_name: str,
    count_min=None,
    count_max=None,
) -> str:
    if not prompt_template:
        prompt_template = (
            "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n"
            "要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"
        )

    if count_min is None or count_max is None:
        default_min, default_max = determine_word_count_bounds(original_count)
        count_min = count_min if count_min is not None else default_min
        count_max = count_max if count_max is not None else default_max

    attempts = 0
    max_attempts = 4
    last_response_text = ""
    while attempts < max_attempts:
        attempts += 1
        if "{text}" in prompt_template:
            current_prompt = prompt_template.format(text=full_text)
        else:
            current_prompt = prompt_template + "\n\n" + full_text

        current_prompt += f"\n\n请注意：这一次的修改后的正文总字数应控制在 {count_min} 到 {count_max} 之间"

        if attempts > 1:
            current_prompt += (
                "\n\n字数不符合规则，请重新修改并返回修改后的正文。"
                " 只输出正文，不要解释。"
                f" 这次要求字数在 {count_min} 到 {count_max} 之间。"
            )

        append_log_ai(f"  🧠 {file_name} AI 第{attempts}次输出，正在检查字数...")

        # 后台日志：打印发给 LLM 的内容到控制台
        print(f"\n[LLM Prompt - Attempt {attempts}] Sending to AI LLM for {file_name}:\n{current_prompt}\n", flush=True)

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "你是一名严谨的中文校对助手"},
                {"role": "user", "content": current_prompt},
            ],
            temperature=0.1,
            stream=False
        )
        result_text = response.choices[0].message.content.strip()
        last_response_text = result_text

        if is_word_count_within_rules(result_text, original_count) if count_min is None and count_max is None else count_min <= count_chinese_characters(result_text) <= count_max:
            count = count_chinese_characters(result_text)
            append_log_ai(
                f"  ✅ {file_name} 字数符合：{count}（目标 {count_min}-{count_max}）"
            )
            return result_text

        current_count = count_chinese_characters(result_text)
        append_log_ai(
            f"  ⚠️ {file_name} 字数不合规：{current_count}，目标 {count_min}-{count_max}，正在重试..."
        )

    raise RuntimeError(
        f"{file_name} 的 AI 输出未能在 {count_min}-{count_max} 范围内，最后一次字数 {current_count}"
    )


def _process_ai(folder, api_key, base_url, prompt_template, process_only_modified=False, count_min=None, count_max=None):
    """步骤 AI：AI 处理 - 保留原始内容，追加修改后的内容"""
    if not prompt_template:
        prompt_template = (
            "下面是一篇中文文章，请你【只修改错别字和明显的识别错误】。\n"
            "要求：1. 不改变原意 2. 不润色文风 3. 不增删内容 4. 保持原有段落结构 5. 只输出修改后的完整文章正文\n"
        )
    
    client = OpenAI(api_key=api_key, base_url=base_url)
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    all_text = "\n".join(
                        [
                            p.text
                            for p in doc.paragraphs
                            if p.text.strip() and p.text.strip() not in ("修改前：", "修改后：")
                        ]
                    )

                    if not all_text.strip():
                        append_log_ai(f"  ⊘ {file} (空文档)")
                        continue

                    original_count = count_chinese_characters(all_text)
                    ai_result = llm_edit_with_word_count_supervision(
                        prompt_template,
                        all_text,
                        client,
                        "deepseek-chat",
                        original_count,
                        file,
                        count_min=count_min,
                        count_max=count_max,
                    )

                    # 在原始内容末尾添加 "修改后：" 标签和分页符
                    last_para = doc.paragraphs[-1] if doc.paragraphs else None
                    if last_para:
                        if last_para.runs:
                            last_para.runs[-1].add_break(WD_BREAK.PAGE)
                        else:
                            last_para.add_run().add_break(WD_BREAK.PAGE)

                    para_modify_label = doc.add_paragraph("修改后：")
                    para_modify_label.paragraph_format.first_line_indent = Cm(0.74)
                    para_modify_label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                    para_modify_label.paragraph_format.line_spacing = Pt(12)

                    for line in ai_result.split("\n"):
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            fmt = p.paragraph_format
                            fmt.first_line_indent = Cm(0.74)
                            fmt.space_before = Pt(0)
                            fmt.space_after = Pt(0)
                            fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            fmt.line_spacing = Pt(12)

                    doc.save(doc_path)
                    append_log_ai(f"  ✓ {file}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")

def _add_labels(folder, process_only_modified=False):
    """步骤 2：添加标签 - 仅当 AI 步骤未运行时才添加分页符和修改后标签"""
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    if doc.paragraphs:
                        # 检查是否已经有"修改后："标签（说明 AI 步骤已执行）
                        last_para = doc.paragraphs[-1]
                        has_modify_label = last_para.text.strip() == "修改后：" or \
                                         (len(doc.paragraphs) > 1 and doc.paragraphs[-2].text.strip() == "修改后：")
                        
                        # 如果还没有"修改前："标签，则添加
                        if doc.paragraphs[0].text.strip() != "修改前：":
                            doc.paragraphs[0].insert_paragraph_before("修改前：")
                        
                        # 如果还没有"修改后："标签（AI 步骤未运行），则添加分页符和修改后标签
                        if not has_modify_label:
                            last_para = doc.paragraphs[-1]
                            if last_para.runs:
                                last_para.runs[-1].add_break(WD_BREAK.PAGE)
                            else:
                                last_para.add_run().add_break(WD_BREAK.PAGE)
                            para_after = doc.add_paragraph("修改后：")
                            para_after.paragraph_format.first_line_indent = Cm(0.74)
                            para_after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                            para_after.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    append_log_ai(f"  ✓ {file}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")

def _format_docs(folder, process_only_modified=False):
    """步骤 3：格式化"""
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    style = doc.styles["Normal"]
                    style.font.name = "宋体"
                    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    style.font.size = Pt(12)
                    for para in doc.paragraphs:
                        para.paragraph_format.first_line_indent = Cm(0.74)
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                        para.paragraph_format.line_spacing = Pt(12)
                    doc.save(doc_path)
                    append_log_ai(f"  ✓ {file}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")

def _set_author(folder, process_only_modified=False):
    """步骤 5：修改作者"""
    for root, files in iter_files_limited(folder, max_depth=4):
        for file in files:
            name_check = file.lstrip()
            name_lower = file.lower()
            # 如果指定只处理修改文件，则跳过不以"改 "开头的文件
            if process_only_modified and not name_check.startswith("改 "):
                continue
            if name_lower.endswith(".docx") and not name_check.startswith("~$"):
                doc_path = os.path.join(root, file)
                try:
                    doc = Document(doc_path)
                    doc.core_properties.author = "思睿教育_美丽可爱的尹老师"
                    doc.save(doc_path)
                    append_log_ai(f"  ✓ {file}")
                except Exception as e:
                    append_log_ai(f"  ✗ {file}: {e}")

ctk.CTkButton(page2, text="🚀 开始流程", fg_color="#2196F3", text_color="white", height=40, command=start_ai_workflow).pack(pady=12)

# 日志
ctk.CTkLabel(page2, text="处理日志", anchor="w").pack(anchor="w", padx=12, pady=(0, 4))
ai_log_text = ctk.CTkTextbox(page2, height=200)
ai_log_text.configure(state="disabled")
ai_log_text.pack(padx=12, pady=(0, 10), fill="both", expand=True)

# 默认显示第一页
show_page("ocr")

root.mainloop()
